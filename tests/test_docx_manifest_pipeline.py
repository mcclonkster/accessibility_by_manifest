import json
from pathlib import Path

from accessibility_by_manifest.inputs.docx.config import DocxManifestConfig
from accessibility_by_manifest.inputs.docx.paths import DocxRun, discover_docx_files, output_paths, plan_runs
from accessibility_by_manifest.pipelines.docx import run_docx


def make_docx(path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = "Accessible DOCX"
    document.add_heading("Accessible DOCX heading", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("This is body text in the DOCX.")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Account"
    table.cell(0, 1).text = "Amount"
    document.save(path)


def test_docx_paths_and_discovery(tmp_path: Path) -> None:
    docx_path = tmp_path / "report.docx"
    lock_path = tmp_path / "~$report.docx"
    docx_path.write_text("not a real docx", encoding="utf-8")
    lock_path.write_text("lock", encoding="utf-8")

    assert discover_docx_files(tmp_path, recursive=False) == [docx_path]
    run = plan_runs([docx_path], docx_path, tmp_path / "out")[0]
    paths = output_paths(run)

    assert paths.manifest_json == tmp_path / "out" / "report_docx_manifest.json"
    assert paths.normalized_manifest_json == tmp_path / "out" / "report_docx_normalized_ir.json"
    assert paths.log_file == tmp_path / "out" / "report.log"


def test_docx_manifest_pipeline_writes_manifest(tmp_path: Path) -> None:
    docx_path = tmp_path / "report.docx"
    output_dir = tmp_path / "out"
    make_docx(docx_path)
    config = DocxManifestConfig(input_path=docx_path, output_root=output_dir, overwrite=True)

    result = run_docx(config, DocxRun(docx_path=docx_path, output_dir=output_dir))

    assert result.ok, result.message
    assert result.output_paths.log_file.exists()
    assert result.output_paths.manifest_json.exists()
    assert result.output_paths.normalized_manifest_json.exists()
    manifest = json.loads(result.output_paths.manifest_json.read_text(encoding="utf-8"))
    normalized_ir = json.loads(result.output_paths.normalized_manifest_json.read_text(encoding="utf-8"))
    assert manifest["manifest_kind"] == "docx_accessibility_manifest"
    assert manifest["source_package"]["format_family"] == "DOCX"
    assert manifest["document_metadata"]["title"] == "Accessible DOCX"
    assert manifest["document_summary"]["paragraph_count"] >= 2
    assert manifest["document_summary"]["body_paragraph_count"] >= 2
    assert manifest["document_summary"]["table_count"] == 1
    assert manifest["document_summary"]["body_table_count"] == 1
    assert manifest["document_summary"]["table_cell_count"] == 2
    assert manifest["document_summary"]["package_xml_paragraph_tag_count"] >= manifest["document_summary"]["body_paragraph_count"]
    assert manifest["document_summary"]["package_xml_table_tag_count"] >= manifest["document_summary"]["body_table_count"]
    assert manifest["document_summary"]["relationship_count"] > 0
    assert manifest["document_summary"]["body_block_count"] >= 3
    assert manifest["document_summary"]["story_part_count"] >= 1
    assert manifest["document_summary"]["raw_style_count"] >= manifest["document_summary"]["style_count"]
    assert manifest["document_summary"]["font_table_entry_count"] >= 1
    assert manifest["document_summary"]["font_family_reference_count"] >= 1
    assert manifest["document_summary"]["direct_run_font_reference_count"] >= 1
    assert manifest["document_summary"]["rebuild_payload_part_count"] == manifest["document_summary"]["package_part_count"]
    assert manifest["document_summary"]["rebuild_payloads_included"] is False
    assert manifest["document_summary"]["raw_xml_included"] is False
    assert manifest["document_summary"]["rebuild_text_payload_part_count"] >= 1
    assert manifest["document_summary"]["rebuild_payload_byte_count"] > 0
    assert manifest["document_summary"]["package_xml_run_tag_count"] >= manifest["document_summary"]["run_count"]
    assert manifest["package_evidence"]["part_kind_counts"]["main_document"] == 1
    document_payload = next(
        entry for entry in manifest["package_rebuild_evidence"]["part_payload_entries"] if entry["part_name"] == "word/document.xml"
    )
    assert document_payload["payload_encoding"] == "utf-8"
    assert document_payload["payload_included"] is False
    assert document_payload["sha256"]
    assert "text" not in document_payload
    assert manifest["body_block_entries"][0]["block_type"] == "paragraph"
    assert manifest["body_block_entries"][0]["raw_xml_sha256"]
    assert "raw_xml" not in manifest["body_block_entries"][0]
    assert "table" in {entry["block_type"] for entry in manifest["body_block_entries"]}
    assert manifest["story_entries"][0]["story_kind"] == "main"
    assert manifest["style_evidence"]["style_count"] >= 1
    assert manifest["style_evidence"]["document_default_language"] is not None
    assert manifest["font_evidence"]["font_table_entries"]
    assert any(entry["font_name"] == "Arial" for entry in manifest["font_evidence"]["font_family_counts"])
    assert "styles" in manifest["relationship_type_counts"]
    assert manifest["paragraph_entries"][0]["text"] == "Accessible DOCX heading"
    assert manifest["paragraph_entries"][1]["runs"][0]["raw_xml_sha256"]
    assert "raw_xml" not in manifest["paragraph_entries"][1]["runs"][0]
    assert manifest["table_entries"][0]["column_count"] == 2
    assert normalized_ir["view_kind"] == "normalized_document_ir"
    assert normalized_ir["source_format"] == "DOCX"
    assert normalized_ir["summary"]["node_count"] >= 3
    assert normalized_ir["summary"]["table_count"] == 1
    assert normalized_ir["summary"]["node_type_counts"]["heading"] >= 1
    assert normalized_ir["summary"]["node_type_counts"]["table"] == 1
    assert any(node["node_type"] == "paragraph" for node in normalized_ir["node_entries"])
    assert normalized_ir["table_entries"][0]["header_candidate_row_indexes"] == [1]
    assert normalized_ir["node_entries"][0]["source_refs"]


def test_docx_manifest_can_inline_rebuild_payloads_when_requested(tmp_path: Path) -> None:
    docx_path = tmp_path / "report.docx"
    output_dir = tmp_path / "out"
    make_docx(docx_path)
    config = DocxManifestConfig(
        input_path=docx_path,
        output_root=output_dir,
        overwrite=True,
        include_rebuild_payloads=True,
    )

    result = run_docx(config, DocxRun(docx_path=docx_path, output_dir=output_dir))

    assert result.ok, result.message
    manifest = json.loads(result.output_paths.manifest_json.read_text(encoding="utf-8"))
    document_payload = next(
        entry for entry in manifest["package_rebuild_evidence"]["part_payload_entries"] if entry["part_name"] == "word/document.xml"
    )
    assert manifest["document_summary"]["rebuild_payloads_included"] is True
    assert manifest["document_summary"]["raw_xml_included"] is True
    assert document_payload["payload_included"] is True
    assert "<w:document" in document_payload["text"]
    assert manifest["body_block_entries"][0]["raw_xml"]
    assert manifest["paragraph_entries"][1]["runs"][0]["raw_xml"]
