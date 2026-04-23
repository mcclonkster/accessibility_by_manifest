import json
import sys
import types
from pathlib import Path
from zipfile import ZipFile

from jsonschema import Draft202012Validator

from accessibility_by_manifest.inputs.pdf.config import PdfManifestConfig
from accessibility_by_manifest.inputs.pdf.paths import PdfRun
from accessibility_by_manifest.cli.pdf import build_parser, config_from_args
from accessibility_by_manifest.manifest.pdf_validate import load_schema
from accessibility_by_manifest.outputs.manifest import PDF_EXTRACTOR_NAMES
from accessibility_by_manifest.pipelines.pdf import run_pdf


def make_pdf(path: Path) -> None:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Accessible PDF heading")
    page.insert_text((72, 100), "This is body text extracted by PyMuPDF.")
    page.insert_link({"kind": fitz.LINK_URI, "from": fitz.Rect(72, 120, 180, 140), "uri": "https://example.com"})
    document.set_metadata({"title": "Accessible PDF", "author": "Test Author"})
    document.save(path)
    document.close()


def append_blank_page(path: Path) -> None:
    import fitz

    document = fitz.open(path)
    document.new_page()
    document.saveIncr()
    document.close()


def append_image_only_page(path: Path) -> None:
    import fitz

    document = fitz.open(path)
    page = document.new_page()
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 20, 20), False)
    pixmap.clear_with(255)
    page.insert_image(fitz.Rect(72, 72, 180, 180), pixmap=pixmap)
    document.saveIncr()
    document.close()


def test_pdf_cli_accepts_ai_parser_flags(tmp_path: Path) -> None:
    pdf_path = tmp_path / "accessible.pdf"
    pdf_path.write_text("not a real pdf", encoding="utf-8")
    output_dir = tmp_path / "out"

    args = build_parser().parse_args(
        [
            str(pdf_path),
            "--output-dir",
            str(output_dir),
            "--ai-parser",
            "docling",
            "--ai-parser-output-mode",
            "artifacts",
            "--ocr-parser",
            "doctr",
        ]
    )
    config = config_from_args(args)

    assert config.ai_parser == "docling"
    assert config.ai_parser_output_mode == "artifacts"
    assert config.ocr_parser == "doctr"


def test_pymupdf_first_pass_writes_schema_valid_manifest(tmp_path: Path) -> None:
    pdf_path = tmp_path / "accessible.pdf"
    output_dir = tmp_path / "out"
    make_pdf(pdf_path)
    config = PdfManifestConfig(input_path=pdf_path, output_root=output_dir, overwrite=True)
    result = run_pdf(config, PdfRun(pdf_path=pdf_path, output_dir=output_dir))

    assert result.ok, result.message
    assert result.output_paths.log_file.exists()
    assert result.output_paths.manifest_json.exists()
    assert result.output_paths.extractor_manifest_dir.exists()
    assert result.output_paths.normalized_manifest_json.exists()
    assert result.output_paths.review_queue_json.exists()
    assert result.output_paths.projected_docx.exists()
    assert not result.output_paths.ai_parser_output_dir.exists()
    assert not result.output_paths.adobe_reference_comparison_json.exists()
    assert not result.output_paths.adobe_reference_comparison_markdown.exists()
    manifest = result.manifest
    assert manifest is not None
    Draft202012Validator(load_schema()).validate(manifest)
    assert manifest["source_package"]["format_family"] == "PDF"
    assert manifest["document_summary"]["total_pages"] == 1
    assert manifest["document_summary"]["raw_block_count"] >= 1
    assert manifest["document_summary"]["normalized_block_count"] >= 1
    assert manifest["document_summary"]["annotation_count"] >= 1
    assert manifest["document_metadata"]["title"] == "Accessible PDF"
    assert manifest["page_entries"][0]["observed_source"]["text_layer_detected"] is True
    assert manifest["raw_block_entries"][0]["source_evidence_type"] == "untagged_text"
    normalized_block = manifest["normalized_block_entries"][0]
    assert normalized_block["observed_source_refs"]
    assert normalized_block["normalized_workflow"]["reading_order_index"] == 1
    assert normalized_block["normalized_workflow"]["interpreted_role"] in {"heading", "paragraph", "list_item", "unknown"}
    assert any(warning["warning_code"] == "DOCUMENT_LANGUAGE_MISSING" for warning in manifest["document_warning_entries"])
    normalized_view = json.loads(result.output_paths.normalized_manifest_json.read_text(encoding="utf-8"))
    assert normalized_view["view_kind"] == "pdf_normalized_manifest_view"
    assert normalized_view["normalized_block_entries"] == manifest["normalized_block_entries"]
    assert normalized_view["normalized_table_entries"] == manifest["normalized_table_entries"]
    review_queue = json.loads(result.output_paths.review_queue_json.read_text(encoding="utf-8"))
    assert review_queue["view_kind"] == "pdf_review_queue_view"
    assert review_queue["review_entries"] == manifest["review_entries"]
    assert review_queue["document_warning_entries"] == manifest["document_warning_entries"]
    with ZipFile(result.output_paths.projected_docx) as docx_zip:
        document_xml = docx_zip.read("word/document.xml").decode("utf-8")
    assert "Accessible PDF heading" in document_xml
    assert "This is body text extracted by PyMuPDF." in document_xml
    assert "Manual review notes" in document_xml
    log_text = result.output_paths.log_file.read_text(encoding="utf-8")
    assert "Pipeline run started" in log_text
    assert "Running PyMuPDF extractor" in log_text
    assert "PDF manifest build completed" in log_text
    for extractor_name in PDF_EXTRACTOR_NAMES:
        extractor_manifest_path = result.output_paths.extractor_manifest_json(extractor_name)
        assert extractor_manifest_path.exists()
        extractor_manifest = json.loads(extractor_manifest_path.read_text(encoding="utf-8"))
        Draft202012Validator(load_schema()).validate(extractor_manifest)
        assert extractor_manifest["extractor_provenance"]["primary_extractor"] == extractor_name
        assert set(extractor_manifest["extractor_evidence"]).issubset({extractor_name})
        for page_entry in extractor_manifest["page_entries"]:
            assert set(page_entry["extractor_evidence"]).issubset({extractor_name})
        for block_entry in extractor_manifest["raw_block_entries"]:
            assert set(block_entry["extractor_evidence"]) == {extractor_name}


def test_pdf_pipeline_writes_adobe_reference_comparison_when_exports_exist(tmp_path: Path) -> None:
    pdf_path = tmp_path / "accessible.pdf"
    output_dir = tmp_path / "out"
    make_pdf(pdf_path)
    write_adobe_reference_exports(pdf_path)
    config = PdfManifestConfig(input_path=pdf_path, output_root=output_dir, overwrite=True)

    result = run_pdf(config, PdfRun(pdf_path=pdf_path, output_dir=output_dir))

    assert result.ok, result.message
    assert result.output_paths.adobe_reference_comparison_json.exists()
    assert result.output_paths.adobe_reference_comparison_markdown.exists()
    comparison = json.loads(result.output_paths.adobe_reference_comparison_json.read_text(encoding="utf-8"))
    assert comparison["report_kind"] == "adobe_reference_comparison"
    assert comparison["reference_summary"]["any_reference_present"] is True
    assert comparison["references"]["docx"]["present"] is True
    assert comparison["references"]["docx"]["paragraph_count"] >= 1
    assert comparison["references"]["html"]["present"] is True
    assert comparison["references"]["html"]["table_count"] == 1
    assert comparison["references"]["html"]["image_count"] == 1
    assert comparison["references"]["xml"]["present"] is True
    assert comparison["references"]["xml"]["table_tag_count"] == 1
    assert comparison["references"]["xml"]["actual_text_attribute_count"] == 1
    assert comparison["references"]["html_asset_dir"]["asset_count"] == 1
    assert "not canonical ground truth" in comparison["comparison_policy"]
    markdown = result.output_paths.adobe_reference_comparison_markdown.read_text(encoding="utf-8")
    assert "# Adobe Reference Comparison" in markdown
    assert "docx: yes" in markdown


def test_pdf_docling_sidecar_writes_artifacts_and_evidence(monkeypatch, tmp_path: Path) -> None:
    install_fake_docling(monkeypatch)
    pdf_path = tmp_path / "accessible.pdf"
    output_dir = tmp_path / "out"
    make_pdf(pdf_path)
    append_blank_page(pdf_path)
    config = PdfManifestConfig(input_path=pdf_path, output_root=output_dir, overwrite=True, ai_parser="docling")

    result = run_pdf(config, PdfRun(pdf_path=pdf_path, output_dir=output_dir))

    assert result.ok, result.message
    assert result.manifest is not None
    manifest = result.manifest
    Draft202012Validator(load_schema()).validate(manifest)
    docling_evidence = manifest["extractor_evidence"]["docling"]
    assert docling_evidence["canonical_status"] == "evidence_only"
    assert docling_evidence["artifacts_written"] is True
    assert docling_evidence["document_claims"]["table_candidate_count"] == 1
    assert docling_evidence["document_claims"]["heading_candidate_count"] == 1
    assert docling_evidence["materialized_evidence"]["raw_block_count"] >= 2
    assert result.output_paths.ai_parser_dir("docling").exists()
    assert (result.output_paths.ai_parser_dir("docling") / "accessible_docling.md").exists()
    assert (result.output_paths.ai_parser_dir("docling") / "accessible_docling.html").exists()
    assert (result.output_paths.ai_parser_dir("docling") / "accessible_docling.json").exists()
    docling_manifest_path = result.output_paths.extractor_manifest_json("docling")
    assert docling_manifest_path.exists()
    docling_manifest = json.loads(docling_manifest_path.read_text(encoding="utf-8"))
    Draft202012Validator(load_schema()).validate(docling_manifest)
    assert docling_manifest["extractor_provenance"]["primary_extractor"] == "docling"
    assert set(docling_manifest["extractor_evidence"]) == {"docling"}
    assert docling_manifest["raw_block_entries"]
    assert "docling" in docling_manifest["page_entries"][0]["extractor_evidence"]
    normalized_block = manifest["normalized_block_entries"][0]
    assert "docling" not in normalized_block["normalized_workflow"]["role_basis"]
    docling_blocks = [
        block
        for block in manifest["normalized_block_entries"]
        if any("docling" in basis for basis in block["normalized_workflow"]["role_basis"])
    ]
    assert docling_blocks


def test_pdf_docling_sidecar_missing_package_fails_clearly(monkeypatch, tmp_path: Path) -> None:
    from accessibility_by_manifest.inputs.pdf.extractors.ai_parsers import docling as docling_adapter

    def missing_docling(_name: str) -> object:
        raise ImportError("missing")

    monkeypatch.setattr(docling_adapter, "import_module", missing_docling)
    pdf_path = tmp_path / "accessible.pdf"
    output_dir = tmp_path / "out"
    make_pdf(pdf_path)
    config = PdfManifestConfig(input_path=pdf_path, output_root=output_dir, overwrite=True, ai_parser="docling")

    result = run_pdf(config, PdfRun(pdf_path=pdf_path, output_dir=output_dir))

    assert not result.ok
    assert "optional 'docling' package is not installed" in result.message
    assert not result.output_paths.ai_parser_output_dir.exists()
    assert not result.output_paths.manifest_json.exists()


def test_pdf_doctr_ocr_sidecar_materializes_image_only_page_text(monkeypatch, tmp_path: Path) -> None:
    install_fake_doctr(monkeypatch)
    pdf_path = tmp_path / "accessible.pdf"
    output_dir = tmp_path / "out"
    make_pdf(pdf_path)
    append_image_only_page(pdf_path)
    config = PdfManifestConfig(input_path=pdf_path, output_root=output_dir, overwrite=True, ocr_parser="doctr")

    result = run_pdf(config, PdfRun(pdf_path=pdf_path, output_dir=output_dir))

    assert result.ok, result.message
    assert result.manifest is not None
    manifest = result.manifest
    Draft202012Validator(load_schema()).validate(manifest)
    doctr_evidence = manifest["extractor_evidence"]["doctr"]
    assert doctr_evidence["canonical_status"] == "evidence_only"
    assert doctr_evidence["page_count"] == 1
    assert doctr_evidence["raw_block_count"] == 1
    doctr_raw_blocks = [block for block in manifest["raw_block_entries"] if "doctr" in block.get("extractor_evidence", {})]
    assert len(doctr_raw_blocks) == 1
    assert doctr_raw_blocks[0]["text"] == "Scanned total"
    assert manifest["page_entries"][1]["observed_source"]["ocr_text_detected"] is True
    normalized_doctr_blocks = [
        block
        for block in manifest["normalized_block_entries"]
        if any("doctr" in basis for basis in block["normalized_workflow"]["role_basis"])
    ]
    assert len(normalized_doctr_blocks) == 1
    assert normalized_doctr_blocks[0]["normalized_workflow"]["interpreted_role"] == "paragraph"
    doctr_manifest = json.loads(result.output_paths.extractor_manifest_json("doctr").read_text(encoding="utf-8"))
    assert doctr_manifest["raw_block_entries"]


def test_pymupdf_first_pass_dry_run_writes_no_manifest(tmp_path: Path) -> None:
    pdf_path = tmp_path / "accessible.pdf"
    output_dir = tmp_path / "out"
    make_pdf(pdf_path)
    config = PdfManifestConfig(input_path=pdf_path, output_root=output_dir, dry_run=True)
    result = run_pdf(config, PdfRun(pdf_path=pdf_path, output_dir=output_dir))

    assert result.ok, result.message
    assert result.output_paths.log_file.exists()
    assert not result.output_paths.manifest_json.exists()
    assert not result.output_paths.extractor_manifest_dir.exists()
    assert not result.output_paths.normalized_manifest_json.exists()
    assert not result.output_paths.review_queue_json.exists()
    assert not result.output_paths.projected_docx.exists()
    assert result.manifest is not None


def install_fake_docling(monkeypatch) -> None:
    docling_package = types.ModuleType("docling")
    docling_package.__path__ = []
    converter_module = types.ModuleType("docling.document_converter")
    converter_module.DocumentConverter = FakeDocumentConverter
    monkeypatch.setitem(sys.modules, "docling", docling_package)
    monkeypatch.setitem(sys.modules, "docling.document_converter", converter_module)


def install_fake_doctr(monkeypatch) -> None:
    doctr_package = types.ModuleType("doctr")
    doctr_package.__path__ = []
    io_module = types.ModuleType("doctr.io")
    models_module = types.ModuleType("doctr.models")
    io_module.DocumentFile = FakeDocumentFile
    models_module.ocr_predictor = fake_ocr_predictor
    monkeypatch.setitem(sys.modules, "doctr", doctr_package)
    monkeypatch.setitem(sys.modules, "doctr.io", io_module)
    monkeypatch.setitem(sys.modules, "doctr.models", models_module)


class FakeDocumentFile:
    @classmethod
    def from_images(cls, image_paths: list[str]) -> list[str]:
        return image_paths


def fake_ocr_predictor(pretrained: bool = True) -> object:
    assert pretrained is True
    return FakeOcrPredictor()


class FakeOcrPredictor:
    def __call__(self, _document: object) -> object:
        return FakeOcrResult()


class FakeOcrResult:
    def export(self) -> dict:
        return {
            "pages": [
                {
                    "blocks": [
                        {
                            "lines": [
                                {
                                    "words": [
                                        {"value": "Scanned", "confidence": 0.95, "geometry": ((0.1, 0.1), (0.3, 0.14))},
                                        {"value": "total", "confidence": 0.93, "geometry": ((0.32, 0.1), (0.45, 0.14))},
                                    ]
                                }
                            ]
                        }
                    ]
                }
            ]
        }


class FakeDocumentConverter:
    def convert(self, _source: str) -> object:
        return types.SimpleNamespace(document=FakeDoclingDocument())


class FakeDoclingDocument:
    tables = ["table-1"]

    def export_to_markdown(self) -> str:
        return "# Accessible PDF heading\n\n| Account | Amount |\n| --- | --- |\n| Assets | $1 |\n"

    def export_to_html(self) -> str:
        return "<h1>Accessible PDF heading</h1><table><tr><th>Account</th><th>Amount</th></tr></table>"

    def export_to_dict(self) -> dict:
        return {
            "texts": [
                {
                    "self_ref": "#/texts/0",
                    "label": "heading",
                    "text": "Accessible PDF heading",
                    "prov": [{"page_no": 2, "bbox": {"l": 72, "t": 120, "r": 300, "b": 140, "coord_origin": "TOPLEFT"}}],
                }
            ],
            "tables": [
                {
                    "self_ref": "#/tables/0",
                    "label": "table",
                    "prov": [{"page_no": 2, "bbox": {"l": 72, "t": 160, "r": 300, "b": 200, "coord_origin": "TOPLEFT"}}],
                    "data": {
                        "table_cells": [
                            {
                                "bbox": {"l": 72, "t": 160, "r": 180, "b": 180, "coord_origin": "TOPLEFT"},
                                "start_row_offset_idx": 0,
                                "start_col_offset_idx": 0,
                                "text": "Assets",
                            },
                            {
                                "bbox": {"l": 180, "t": 160, "r": 300, "b": 180, "coord_origin": "TOPLEFT"},
                                "start_row_offset_idx": 0,
                                "start_col_offset_idx": 1,
                                "text": "$1",
                            },
                        ]
                    },
                }
            ],
        }


def write_adobe_reference_exports(pdf_path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Accessible PDF heading", level=1)
    document.add_paragraph("This is Adobe reference body text.")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Assets"
    document.save(pdf_path.with_suffix(".docx"))

    html_asset_dir = pdf_path.parent / f"{pdf_path.stem}_files"
    html_asset_dir.mkdir()
    (html_asset_dir / "Image_001.png").write_bytes(b"fake-png")
    pdf_path.with_suffix(".html").write_text(
        '<html><body><h1>Accessible PDF heading</h1><table><tr><td>Assets</td></tr></table>'
        f'<img src="{pdf_path.stem}_files/Image_001.png" alt="Chart"></body></html>',
        encoding="utf-8",
    )
    pdf_path.with_suffix(".xml").write_text(
        '<!-- Created from PDF via Acrobat SaveAsXML --><TaggedPDF-doc><Document>'
        '<P ActualText="Accessible PDF heading">Accessible PDF heading</P>'
        '<Table><TR><TD><P>Assets</P></TD></TR></Table>'
        '<Figure Alt="Chart"><ImageData /></Figure>'
        "</Document></TaggedPDF-doc>",
        encoding="utf-8",
    )
