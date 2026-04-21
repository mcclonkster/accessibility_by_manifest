from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

from accessibility_by_manifest.errors import OutputWriteError
from accessibility_by_manifest.inputs.pdf.paths import PdfOutputPaths


def write_pdf_docx_output(manifest: dict[str, Any], output_paths: PdfOutputPaths, overwrite: bool) -> None:
    document = build_pdf_docx(manifest)
    atomic_write_docx(document, output_paths.projected_docx, overwrite)


def build_pdf_docx(manifest: dict[str, Any]) -> object:
    from docx import Document

    document = Document()
    title = document_title(manifest)
    document.core_properties.title = title
    document.add_heading(title, level=0)
    document.add_paragraph("Draft PDF projection generated from the normalized manifest. Manual review items are preserved below.")

    raw_text_by_normalized_id = {
        f"n_{raw_block['block_id']}": raw_block.get("text", "").strip()
        for raw_block in manifest.get("raw_block_entries", [])
    }
    tables_by_id = {table["table_id"]: table for table in manifest.get("normalized_table_entries", [])}
    rendered_tables: set[str] = set()
    for block in sorted(manifest.get("normalized_block_entries", []), key=reading_order_index):
        role = block.get("normalized_workflow", {}).get("interpreted_role")
        if role == "artifact":
            continue
        if role == "table_cell":
            table_id = block.get("projected_target", {}).get("target_table_id")
            if table_id and table_id not in rendered_tables:
                add_table(document, tables_by_id.get(table_id), raw_text_by_normalized_id)
                rendered_tables.add(table_id)
            elif not table_id:
                add_paragraph_for_block(document, block, raw_text_by_normalized_id)
            continue
        add_paragraph_for_block(document, block, raw_text_by_normalized_id)

    add_review_appendix(document, manifest)
    return document


def document_title(manifest: dict[str, Any]) -> str:
    metadata_title = manifest.get("document_metadata", {}).get("title")
    if metadata_title:
        return f"{metadata_title} draft DOCX"
    file_name = manifest.get("source_package", {}).get("input_file_name") or "PDF"
    return f"{Path(file_name).stem} draft DOCX"


def reading_order_index(block: dict[str, Any]) -> int:
    value = block.get("normalized_workflow", {}).get("reading_order_index")
    return int(value) if isinstance(value, int) else 0


def add_paragraph_for_block(document: object, block: dict[str, Any], raw_text_by_normalized_id: dict[str, str]) -> None:
    text = raw_text_by_normalized_id.get(block.get("block_id"), "").strip()
    if not text:
        return
    workflow = block.get("normalized_workflow", {})
    role = workflow.get("interpreted_role")
    if role == "heading":
        level = workflow.get("heading_level") or 2
        document.add_heading(text, level=max(1, min(int(level), 4)))
    elif role == "list_item":
        style = "List Number" if workflow.get("list_kind") == "number" else "List Bullet"
        document.add_paragraph(text, style=style)
    else:
        document.add_paragraph(text)


def add_table(document: object, table_entry: dict[str, Any] | None, raw_text_by_normalized_id: dict[str, str]) -> None:
    if not table_entry:
        return
    rows = table_entry.get("row_entries", [])
    if not rows:
        return
    column_count = max((len(row.get("cell_block_ids", [])) for row in rows), default=1)
    docx_table = document.add_table(rows=len(rows), cols=max(column_count, 1))
    docx_table.style = "Table Grid"
    header_rows = set(table_entry.get("header_candidate_row_indexes", []))
    for row_index, row_entry in enumerate(rows):
        cell_ids = row_entry.get("cell_block_ids", [])
        for column_index in range(column_count):
            text = raw_text_by_normalized_id.get(cell_ids[column_index], "") if column_index < len(cell_ids) else ""
            cell = docx_table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            if row_entry.get("row_index") in header_rows:
                run.bold = True
    document.add_paragraph()


def add_review_appendix(document: object, manifest: dict[str, Any]) -> None:
    warnings = manifest.get("document_warning_entries", [])
    review_entries = manifest.get("review_entries", [])
    if not warnings and not review_entries:
        return
    document.add_heading("Manual review notes", level=1)
    for warning in warnings:
        document.add_paragraph(
            f"{warning.get('warning_code')}: {warning.get('warning_message')}",
            style="List Bullet",
        )
    for entry in review_entries:
        target = f" [{entry.get('target_ref')}]" if entry.get("target_ref") else ""
        document.add_paragraph(
            f"{entry.get('issue_code')}{target}: {entry.get('reason')}",
            style="List Bullet",
        )


def atomic_write_docx(document: object, path: Path, overwrite: bool) -> None:
    if path.exists() and not overwrite:
        raise OutputWriteError(f"Output already exists. Use --overwrite or choose another folder: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(path.parent), prefix=f"{path.stem}_") as temp_file:
            temp_path = Path(temp_file.name)
        document.save(str(temp_path))
        temp_path.replace(path)
    except Exception as exc:
        if temp_path and temp_path.exists():
            temp_path.unlink(missing_ok=True)
        raise OutputWriteError(f"Failed to write DOCX '{path}': {exc}") from exc
