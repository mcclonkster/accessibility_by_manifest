from __future__ import annotations

import tempfile
from pathlib import Path

from accessibility_by_manifest.errors import OutputWriteError
from accessibility_by_manifest.inputs.pptx.paths import OutputPaths
from accessibility_by_manifest.manifest.pptx import Manifest, ParagraphEntry, SlideEntry
from accessibility_by_manifest.outputs.docx_accessibility import apply_picture_alt_text, enforce_wcag_aaa_text_colors, slide_image_alt_text


SLIDE_IMAGE_HEIGHT_INCHES = 3.75


def write_companion_docx(manifest: Manifest, output_paths: OutputPaths, overwrite: bool) -> None:
    document = build_docx(manifest, title_suffix="companion document", include_warning_notes=False)
    atomic_write_docx(document, output_paths.companion_docx, overwrite)


def write_slides_docx(manifest: Manifest, output_paths: OutputPaths, overwrite: bool) -> None:
    document = build_docx(manifest, title_suffix="slides companion document", include_warning_notes=True)
    atomic_write_docx(document, output_paths.slides_docx, overwrite)


def write_docx_outputs(manifest: Manifest, output_paths: OutputPaths, overwrite: bool) -> None:
    write_companion_docx(manifest, output_paths, overwrite)
    write_slides_docx(manifest, output_paths, overwrite)


def build_docx(manifest: Manifest, title_suffix: str, include_warning_notes: bool) -> object:
    from docx import Document
    from docx.enum.text import WD_BREAK

    title = f"{manifest.source_path.stem} {title_suffix}"
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)

    for index, slide in enumerate(manifest.slides, start=1):
        document.add_heading(f"Slide {slide.slide_number}. {slide.title}", level=1)
        add_slide_image(document, manifest, slide)
        add_slide_text(document, slide)
        add_visual_descriptions(document, slide)
        if include_warning_notes:
            add_warning_notes(document, slide)
        if index < len(manifest.slides):
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    enforce_wcag_aaa_text_colors(document)
    return document


def add_slide_image(document: object, manifest: Manifest, slide: SlideEntry) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    if manifest.preview_image_dir is None or slide.preview_image is None:
        document.add_paragraph("[Slide image unavailable. Review required.]")
        return
    image_path = manifest.preview_image_dir / slide.preview_image
    if not image_path.exists():
        document.add_paragraph("[Slide image unavailable. Review required.]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline_shape = paragraph.add_run().add_picture(str(image_path), height=Inches(SLIDE_IMAGE_HEIGHT_INCHES))
    alt_title, alt_description = slide_image_alt_text(slide.slide_number, slide.title)
    apply_picture_alt_text(inline_shape, alt_title, alt_description)


def add_slide_text(document: object, slide: SlideEntry) -> None:
    document.add_heading("Slide text", level=2)
    paragraphs = [paragraph for block in sorted(slide.text_blocks, key=lambda item: item.order) for paragraph in block.paragraphs]
    if not paragraphs:
        document.add_paragraph("No extractable visible slide text found.")
        return
    for paragraph in paragraphs:
        add_projected_paragraph(document, paragraph)


def add_projected_paragraph(document: object, paragraph: ParagraphEntry) -> None:
    if paragraph.list_kind == "bullet":
        style = "List Bullet" if paragraph.level <= 0 else "List Bullet 2"
        document.add_paragraph(paragraph.text, style=style)
    elif paragraph.list_kind == "number":
        document.add_paragraph(f"[numbered] {paragraph.text}")
    else:
        document.add_paragraph(paragraph.text)


def add_visual_descriptions(document: object, slide: SlideEntry) -> None:
    if not slide.visuals:
        return
    document.add_heading("Visual descriptions", level=2)
    for visual in slide.visuals:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{visual.label}. {visual.visual_type.capitalize()}: ").bold = True
        description = visual.description or "Description metadata missing. Review needed."
        if visual.description_source == "title":
            description += " [metadata source: title]"
        paragraph.add_run(description)


def add_warning_notes(document: object, slide: SlideEntry) -> None:
    warnings = list(slide.warnings)
    for block in slide.text_blocks:
        warnings.extend(block.warnings)
        for paragraph in block.paragraphs:
            warnings.extend(paragraph.warnings)
    for visual in slide.visuals:
        warnings.extend(visual.warnings)
    if not warnings:
        return
    document.add_heading("Manual review notes", level=2)
    for warning in warnings:
        document.add_paragraph(f"{warning.code}: {warning.message}", style="List Bullet")


def atomic_write_docx(document: object, path: Path, overwrite: bool) -> None:
    if path.exists() and not overwrite:
        raise OutputWriteError(f"Output already exists. Use --overwrite or choose another folder: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(path.parent), prefix=f"{path.stem}_") as temp_file:
            temp_path = Path(temp_file.name)
        document.save(str(temp_path))
        temp_path.replace(path)
    except Exception as exc:
        if temp_path and temp_path.exists():
            temp_path.unlink(missing_ok=True)
        raise OutputWriteError(f"Failed to write DOCX '{path}': {exc}") from exc
