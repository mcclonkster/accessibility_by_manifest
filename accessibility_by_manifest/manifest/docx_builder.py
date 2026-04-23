from __future__ import annotations

import base64
import hashlib
import re
from collections import Counter
from datetime import date, datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from accessibility_by_manifest.inputs.docx.paths import DocxOutputPaths
from accessibility_by_manifest.util.fingerprints import file_sha256
from accessibility_by_manifest.util.logging import get_logger


LOGGER = get_logger(__name__)


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}

def build_docx_manifest(
    docx_path: Path,
    output_paths: DocxOutputPaths,
    *,
    include_rebuild_payloads: bool = False,
) -> dict[str, Any]:
    LOGGER.info("Building DOCX manifest for %s", docx_path)
    from docx import Document

    document = Document(docx_path)
    package, xml_parts, package_rebuild_evidence = inspect_docx_package(
        docx_path,
        include_rebuild_payloads=include_rebuild_payloads,
    )
    package_xml_summary = summarize_package_xml(package)
    body_block_entries = body_block_entries_from_document_xml(
        xml_parts.get("word/document.xml", ""),
        include_raw_xml=include_rebuild_payloads,
    )
    story_entries = story_entries_from_xml_parts(xml_parts)
    drawing_entries = drawing_entries_from_xml_parts(xml_parts)
    hyperlink_entries = hyperlink_entries_from_xml_parts(xml_parts)
    paragraph_entries = paragraph_entries_from_document(document, include_raw_xml=include_rebuild_payloads)
    table_entries, table_cell_entries = table_entries_from_document(document)
    section_entries = section_entries_from_document(document)
    styles = style_entries_from_document(document)
    style_evidence = style_evidence_from_xml(xml_parts.get("word/styles.xml", ""))
    numbering_evidence = numbering_evidence_from_xml(xml_parts.get("word/numbering.xml", ""))
    settings_evidence = settings_evidence_from_xml(xml_parts.get("word/settings.xml", ""))
    extended_properties = extended_properties_from_xml(xml_parts.get("docProps/app.xml", ""))
    font_evidence = font_evidence_from_xml_parts(xml_parts)
    media_entries = media_entries_from_package(package)
    header_footer_entries = header_footer_entries_from_package(package)
    relationship_entries = relationship_entries_from_package(package)
    relationship_type_counts = relationship_type_counts_from_entries(relationship_entries)
    summary = {
        "paragraph_count": len(paragraph_entries),
        "body_paragraph_count": len(paragraph_entries),
        "table_count": len(table_entries),
        "body_table_count": len(table_entries),
        "body_block_count": len(body_block_entries),
        "table_cell_count": len(table_cell_entries),
        "run_count": sum(entry["run_count"] for entry in paragraph_entries),
        "hyperlink_count": sum(entry["hyperlink_count"] for entry in paragraph_entries),
        "wordprocessingml_hyperlink_count": len(hyperlink_entries),
        "numbered_paragraph_count": sum(1 for entry in paragraph_entries if entry["numbering_detected"]),
        "wordprocessingml_numbering_reference_count": sum(1 for entry in body_block_entries if entry.get("numbering")),
        "image_count": len(media_entries),
        "drawing_count": len(drawing_entries),
        "drawing_with_alt_text_count": sum(1 for entry in drawing_entries if entry.get("title") or entry.get("description")),
        "header_count": sum(1 for entry in header_footer_entries if entry["part_kind"] == "header"),
        "footer_count": sum(1 for entry in header_footer_entries if entry["part_kind"] == "footer"),
        "story_part_count": len(story_entries),
        "section_count": len(section_entries),
        "style_count": len(styles),
        "raw_style_count": style_evidence["style_count"],
        "font_table_entry_count": len(font_evidence["font_table_entries"]),
        "font_family_reference_count": len(font_evidence["font_family_counts"]),
        "style_font_reference_count": len(font_evidence["style_font_entries"]),
        "direct_run_font_reference_count": len(font_evidence["direct_run_font_entries"]),
        "theme_font_reference_count": len(font_evidence["theme_font_entries"]),
        "numbering_definition_count": numbering_evidence["numbering_definition_count"],
        "package_part_count": len(package["part_entries"]),
        "rebuild_payload_part_count": len(package_rebuild_evidence["part_payload_entries"]),
        "rebuild_payloads_included": package_rebuild_evidence["payloads_included"],
        "raw_xml_included": include_rebuild_payloads,
        "rebuild_text_payload_part_count": package_rebuild_evidence["text_payload_part_count"],
        "rebuild_binary_payload_part_count": package_rebuild_evidence["binary_payload_part_count"],
        "rebuild_payload_byte_count": package_rebuild_evidence["payload_byte_count"],
        "relationship_count": len(relationship_entries),
        "external_relationship_count": sum(1 for entry in relationship_entries if entry.get("target_mode") == "External"),
        **package_xml_summary,
    }
    return {
        "manifest_kind": "docx_accessibility_manifest",
        "manifest_version": "0.1",
        "source_package": {
            "format_family": "DOCX",
            "input_file_name": docx_path.name,
            "input_file_path": str(docx_path),
            "byte_size": docx_path.stat().st_size,
            "sha256": file_sha256(docx_path),
        },
        "target_package": {
            "manifest_json_file_name": output_paths.manifest_json.name,
            "manifest_json_file_path": str(output_paths.manifest_json),
        },
        "document_metadata": core_properties(document),
        "extended_properties": extended_properties,
        "document_summary": summary,
        "package_evidence": package,
        "package_rebuild_evidence": package_rebuild_evidence,
        "relationship_entries": relationship_entries,
        "relationship_type_counts": relationship_type_counts,
        "settings_evidence": settings_evidence,
        "style_evidence": style_evidence,
        "font_evidence": font_evidence,
        "numbering_evidence": numbering_evidence,
        "style_entries": styles,
        "section_entries": section_entries,
        "story_entries": story_entries,
        "header_footer_entries": header_footer_entries,
        "media_entries": media_entries,
        "drawing_entries": drawing_entries,
        "hyperlink_entries": hyperlink_entries,
        "body_block_entries": body_block_entries,
        "paragraph_entries": paragraph_entries,
        "table_entries": table_entries,
        "table_cell_entries": table_cell_entries,
        "document_warning_entries": document_warnings(document, summary, style_evidence),
        "notes": [
            "This is a DOCX input manifest. It records Adobe's DOCX export as a source document, not as canonical truth for the original PDF.",
            "The manifest preserves both the python-docx body model and lower-level WordprocessingML/package evidence.",
            "package_rebuild_evidence stores part metadata and checksums by default; raw OPC package payloads are only inlined when include_rebuild_payloads is enabled.",
            "Paragraph, table, media, relationship, story, and package evidence can be compared against the PDF master manifest and Adobe reference comparison outputs.",
        ],
    }


def inspect_docx_package(
    docx_path: Path,
    *,
    include_rebuild_payloads: bool = False,
) -> tuple[dict[str, Any], dict[str, str], dict[str, Any]]:
    with ZipFile(docx_path) as docx_zip:
        names = sorted(docx_zip.namelist())
        part_entries = []
        payload_entries = []
        xml_parts = {}
        for name in names:
            if name.endswith("/"):
                continue
            info = docx_zip.getinfo(name)
            payload = docx_zip.read(name)
            part_entry = {
                "part_name": name,
                "part_kind": classify_part_name(name),
                "byte_size": info.file_size,
                "compressed_byte_size": info.compress_size,
                "suffix": Path(name).suffix.lower(),
                "sha256": bytes_sha256(payload),
                "crc": info.CRC,
                "zip_compress_type": info.compress_type,
                "zip_date_time": list(info.date_time),
            }
            part_entries.append(part_entry)
            payload_entry = {
                **part_entry,
                "payload_encoding": payload_encoding_for_part(name),
                "payload_included": include_rebuild_payloads,
            }
            if payload_entry["payload_encoding"] == "utf-8":
                text = payload.decode("utf-8", errors="replace")
                xml_parts[name] = text
                if include_rebuild_payloads:
                    payload_entry["text"] = text
                    payload_entry["base64"] = None
            else:
                if include_rebuild_payloads:
                    payload_entry["text"] = None
                    payload_entry["base64"] = base64.b64encode(payload).decode("ascii")
            payload_entries.append(payload_entry)
    content_types = parse_content_types(xml_parts.get("[Content_Types].xml", ""))
    package = {
        "part_entries": part_entries,
        "part_kind_counts": part_kind_counts(part_entries),
        "content_types": content_types,
        "relationship_entries": parse_relationship_entries(xml_parts),
        "xml_part_metrics": [
            {
                "part_name": name,
                "part_kind": classify_part_name(name),
                "paragraph_tag_count": count_regex(text, r"<w:p(?:\s|>)"),
                "run_tag_count": count_regex(text, r"<w:r(?:\s|>)"),
                "table_tag_count": count_regex(text, r"<w:tbl(?:\s|>)"),
                "table_row_tag_count": count_regex(text, r"<w:tr(?:\s|>)"),
                "table_cell_tag_count": count_regex(text, r"<w:tc(?:\s|>)"),
                "drawing_tag_count": count_regex(text, r"<w:drawing(?:\s|>)"),
                "pict_tag_count": count_regex(text, r"<w:pict(?:\s|>)"),
                "hyperlink_tag_count": count_regex(text, r"<w:hyperlink(?:\s|>)"),
                "text_tag_count": count_regex(text, r"<w:t(?:\s|>)"),
                "section_properties_count": count_regex(text, r"<w:sectPr(?:\s|>)"),
                "break_tag_count": count_regex(text, r"<w:br(?:\s|>)"),
                "last_rendered_page_break_count": count_regex(text, r"<w:lastRenderedPageBreak(?:\s|/>)"),
                "bookmark_start_count": count_regex(text, r"<w:bookmarkStart(?:\s|/>)"),
                "comment_reference_count": count_regex(text, r"<w:commentReference(?:\s|/>)"),
                "footnote_reference_count": count_regex(text, r"<w:footnoteReference(?:\s|/>)"),
                "endnote_reference_count": count_regex(text, r"<w:endnoteReference(?:\s|/>)"),
                "field_char_count": count_regex(text, r"<w:fldChar(?:\s|>)"),
                "instruction_text_count": count_regex(text, r"<w:instrText(?:\s|>)"),
                "language_tag_count": count_regex(text, r"<w:lang(?:\s|/>)"),
                "docpr_count": count_regex(text, r"<wp:docPr(?:\s|/>)"),
                "docpr_with_title_count": count_regex(text, r"<wp:docPr\b[^>]*\btitle="),
                "docpr_with_description_count": count_regex(text, r"<wp:docPr\b[^>]*\bdescr="),
            }
            for name, text in sorted(xml_parts.items())
        ],
    }
    rebuild_evidence = {
        "rebuild_strategy": "opc_package_part_replay",
        "rebuild_scope": "all_non_directory_zip_parts",
        "payloads_included": include_rebuild_payloads,
        "part_payload_entries": payload_entries,
        "part_order": [entry["part_name"] for entry in payload_entries],
        "text_payload_part_count": sum(1 for entry in payload_entries if entry["payload_encoding"] == "utf-8"),
        "binary_payload_part_count": sum(1 for entry in payload_entries if entry["payload_encoding"] == "base64"),
        "payload_byte_count": sum(entry["byte_size"] for entry in payload_entries),
        "notes": [
            "This section preserves package part metadata and checksums for recreation planning.",
            "When payloads_included is false, payload entries are references/provenance records, not embedded source bytes.",
            "When payloads_included is true, XML and relationship parts are stored as UTF-8 text, and binary package parts are stored as base64.",
        ],
    }
    return package, xml_parts, rebuild_evidence


def payload_encoding_for_part(part_name: str) -> str:
    if part_name.endswith(".xml") or part_name.endswith(".rels"):
        return "utf-8"
    return "base64"


def bytes_sha256(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def summarize_package_xml(package: dict[str, Any]) -> dict[str, int]:
    metrics = package["xml_part_metrics"]
    return {
        "package_xml_paragraph_tag_count": sum(entry["paragraph_tag_count"] for entry in metrics),
        "package_xml_run_tag_count": sum(entry["run_tag_count"] for entry in metrics),
        "package_xml_table_tag_count": sum(entry["table_tag_count"] for entry in metrics),
        "package_xml_table_row_tag_count": sum(entry["table_row_tag_count"] for entry in metrics),
        "package_xml_table_cell_tag_count": sum(entry["table_cell_tag_count"] for entry in metrics),
        "package_xml_drawing_tag_count": sum(entry["drawing_tag_count"] for entry in metrics),
        "package_xml_pict_tag_count": sum(entry["pict_tag_count"] for entry in metrics),
        "package_xml_hyperlink_tag_count": sum(entry["hyperlink_tag_count"] for entry in metrics),
        "package_xml_text_tag_count": sum(entry["text_tag_count"] for entry in metrics),
        "package_xml_section_properties_count": sum(entry["section_properties_count"] for entry in metrics),
        "package_xml_bookmark_start_count": sum(entry["bookmark_start_count"] for entry in metrics),
        "package_xml_comment_reference_count": sum(entry["comment_reference_count"] for entry in metrics),
        "package_xml_footnote_reference_count": sum(entry["footnote_reference_count"] for entry in metrics),
        "package_xml_endnote_reference_count": sum(entry["endnote_reference_count"] for entry in metrics),
        "package_xml_field_char_count": sum(entry["field_char_count"] for entry in metrics),
        "package_xml_language_tag_count": sum(entry["language_tag_count"] for entry in metrics),
        "package_xml_docpr_count": sum(entry["docpr_count"] for entry in metrics),
        "package_xml_docpr_with_alt_text_count": sum(
            entry["docpr_with_title_count"] + entry["docpr_with_description_count"] for entry in metrics
        ),
    }


def parse_content_types(content_types_xml: str) -> dict[str, Any]:
    defaults = re.findall(r'<Default\s+[^>]*Extension="([^"]+)"[^>]*ContentType="([^"]+)"', content_types_xml)
    overrides = re.findall(r'<Override\s+[^>]*PartName="([^"]+)"[^>]*ContentType="([^"]+)"', content_types_xml)
    return {
        "default_count": len(defaults),
        "override_count": len(overrides),
        "defaults": [{"extension": extension, "content_type": content_type} for extension, content_type in defaults],
        "overrides": [{"part_name": part_name, "content_type": content_type} for part_name, content_type in overrides],
    }


def part_kind_counts(part_entries: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for part in part_entries:
        kind = classify_part_name(part["part_name"])
        counts[kind] = counts.get(kind, 0) + 1
    return counts


def classify_part_name(part_name: str) -> str:
    if part_name == "word/document.xml":
        return "main_document"
    if part_name.startswith("word/header"):
        return "header"
    if part_name.startswith("word/footer"):
        return "footer"
    if part_name == "word/styles.xml":
        return "styles"
    if part_name == "word/numbering.xml":
        return "numbering"
    if part_name == "word/settings.xml":
        return "settings"
    if part_name == "word/fontTable.xml":
        return "font_table"
    if part_name == "word/comments.xml":
        return "comments"
    if part_name == "word/footnotes.xml":
        return "footnotes"
    if part_name == "word/endnotes.xml":
        return "endnotes"
    if part_name.startswith("word/media/"):
        return "media"
    if part_name.startswith("word/theme/"):
        return "theme"
    if part_name.startswith("docProps/"):
        return "document_properties"
    if part_name.startswith("customXml/"):
        return "custom_xml"
    if part_name.endswith(".rels"):
        return "relationships"
    return "other"


def relationship_entries_from_package(package: dict[str, Any]) -> list[dict[str, Any]]:
    return list(package.get("relationship_entries", []))


def parse_relationship_entries(xml_parts: dict[str, str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for part_name, text in sorted(xml_parts.items()):
        if not part_name.endswith(".rels"):
            continue
        for relationship_xml in re.findall(r"<Relationship\b[^>]*/?>", text):
            attrs = xml_attributes(relationship_xml)
            entries.append(
                {
                    "relationship_part": part_name,
                    "relationship_id": attrs.get("Id"),
                    "type": attrs.get("Type"),
                    "type_name": relationship_type_name(attrs.get("Type")),
                    "target": attrs.get("Target"),
                    "target_mode": attrs.get("TargetMode"),
                }
            )
    return entries


def xml_attributes(xml: str) -> dict[str, str]:
    return dict(re.findall(r"([A-Za-z_:][A-Za-z0-9_.:-]*)=\"([^\"]*)\"", xml))


def relationship_type_name(type_uri: str | None) -> str | None:
    if not type_uri:
        return None
    return type_uri.rstrip("/").rsplit("/", 1)[-1]


def relationship_type_counts_from_entries(entries: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for entry in entries:
        type_name = entry.get("type_name") or "unknown"
        counts[type_name] = counts.get(type_name, 0) + 1
    return counts


def body_block_entries_from_document_xml(document_xml: str, *, include_raw_xml: bool = False) -> list[dict[str, Any]]:
    root = parse_xml(document_xml)
    if root is None:
        return []
    body = root.find("w:body", NS)
    if body is None:
        return []

    entries: list[dict[str, Any]] = []
    block_index = 1
    for child in list(body):
        if child.tag == qn("w", "p"):
            entries.append(paragraph_block_entry(child, block_index, include_raw_xml=include_raw_xml))
            block_index += 1
        elif child.tag == qn("w", "tbl"):
            entries.append(table_block_entry(child, block_index, include_raw_xml=include_raw_xml))
            block_index += 1
        elif child.tag == qn("w", "sectPr"):
            entries.append(section_properties_block_entry(child, block_index, include_raw_xml=include_raw_xml))
            block_index += 1
    return entries


def paragraph_block_entry(paragraph: ET.Element, block_index: int, *, include_raw_xml: bool = False) -> dict[str, Any]:
    ppr = paragraph.find("w:pPr", NS)
    style = ppr.find("w:pStyle", NS) if ppr is not None else None
    outline = ppr.find("w:outlineLvl", NS) if ppr is not None else None
    justification = ppr.find("w:jc", NS) if ppr is not None else None
    raw_xml = element_xml(paragraph)
    paragraph_properties_xml = element_xml(ppr)
    entry = {
        "block_index": block_index,
        "block_type": "paragraph",
        "text": element_text(paragraph),
        "style_id": attr(style, "w", "val"),
        "outline_level": attr(outline, "w", "val"),
        "justification": attr(justification, "w", "val"),
        "numbering": paragraph_numbering(paragraph),
        "paragraph_properties_present": ppr is not None,
        "paragraph_properties_xml_sha256": text_sha256(paragraph_properties_xml),
        "raw_xml_sha256": text_sha256(raw_xml),
        "run_count": len(paragraph.findall("w:r", NS)),
        "hyperlink_count": len(paragraph.findall(".//w:hyperlink", NS)),
        "drawing_count": len(paragraph.findall(".//w:drawing", NS)) + len(paragraph.findall(".//w:pict", NS)),
        "bookmark_start_count": len(paragraph.findall(".//w:bookmarkStart", NS)),
        "field_char_count": len(paragraph.findall(".//w:fldChar", NS)),
        "break_count": len(paragraph.findall(".//w:br", NS)),
        "last_rendered_page_break_count": len(paragraph.findall(".//w:lastRenderedPageBreak", NS)),
    }
    if include_raw_xml:
        entry["paragraph_properties_xml"] = paragraph_properties_xml
        entry["raw_xml"] = raw_xml
    return entry


def table_block_entry(table: ET.Element, block_index: int, *, include_raw_xml: bool = False) -> dict[str, Any]:
    rows = table.findall("w:tr", NS)
    cells = table.findall(".//w:tc", NS)
    grid = table.find("w:tblGrid", NS)
    grid_cols = grid.findall("w:gridCol", NS) if grid is not None else []
    table_properties_xml = element_xml(table.find("w:tblPr", NS))
    table_grid_xml = element_xml(grid)
    raw_xml = element_xml(table)
    entry = {
        "block_index": block_index,
        "block_type": "table",
        "row_count": len(rows),
        "cell_count": len(cells),
        "grid_column_count": len(grid_cols),
        "grid_span_count": len(table.findall(".//w:gridSpan", NS)),
        "vertical_merge_count": len(table.findall(".//w:vMerge", NS)),
        "text_sample": table_xml_text_sample(table),
        "table_properties_present": table_properties_xml is not None,
        "table_properties_xml_sha256": text_sha256(table_properties_xml),
        "table_grid_xml_sha256": text_sha256(table_grid_xml),
        "raw_xml_sha256": text_sha256(raw_xml),
    }
    if include_raw_xml:
        entry["table_properties_xml"] = table_properties_xml
        entry["table_grid_xml"] = table_grid_xml
        entry["raw_xml"] = raw_xml
    return entry


def section_properties_block_entry(section_properties: ET.Element, block_index: int, *, include_raw_xml: bool = False) -> dict[str, Any]:
    page_size = section_properties.find("w:pgSz", NS)
    page_margin = section_properties.find("w:pgMar", NS)
    section_type = section_properties.find("w:type", NS)
    raw_xml = element_xml(section_properties)
    entry = {
        "block_index": block_index,
        "block_type": "section_properties",
        "section_type": attr(section_type, "w", "val"),
        "page_width_twips": parse_int_attr(page_size, "w", "w"),
        "page_height_twips": parse_int_attr(page_size, "w", "h"),
        "orientation": attr(page_size, "w", "orient"),
        "top_margin_twips": parse_int_attr(page_margin, "w", "top"),
        "bottom_margin_twips": parse_int_attr(page_margin, "w", "bottom"),
        "left_margin_twips": parse_int_attr(page_margin, "w", "left"),
        "right_margin_twips": parse_int_attr(page_margin, "w", "right"),
        "raw_xml_sha256": text_sha256(raw_xml),
    }
    if include_raw_xml:
        entry["raw_xml"] = raw_xml
    return entry


def paragraph_numbering(paragraph: ET.Element) -> dict[str, str | None] | None:
    num_pr = paragraph.find("w:pPr/w:numPr", NS)
    if num_pr is None:
        return None
    ilvl = num_pr.find("w:ilvl", NS)
    num_id = num_pr.find("w:numId", NS)
    return {
        "level": attr(ilvl, "w", "val"),
        "num_id": attr(num_id, "w", "val"),
    }


def table_xml_text_sample(table: ET.Element) -> list[list[str]]:
    sample: list[list[str]] = []
    for row in table.findall("w:tr", NS)[:5]:
        sample.append([element_text(cell) for cell in row.findall("w:tc", NS)[:8]])
    return sample


def story_entries_from_xml_parts(xml_parts: dict[str, str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for part_name, text in sorted(xml_parts.items()):
        story_kind = story_kind_for_part(part_name)
        if story_kind is None:
            continue
        root = parse_xml(text)
        entries.append(
            {
                "part_name": part_name,
                "story_kind": story_kind,
                "parseable_xml": root is not None,
                "paragraph_count": count_regex(text, r"<w:p(?:\s|>)"),
                "table_count": count_regex(text, r"<w:tbl(?:\s|>)"),
                "drawing_count": count_regex(text, r"<w:drawing(?:\s|>)") + count_regex(text, r"<w:pict(?:\s|>)"),
                "hyperlink_count": count_regex(text, r"<w:hyperlink(?:\s|>)"),
                "text_character_count": len(element_text(root)) if root is not None else None,
                "text_sample": element_text(root)[:500] if root is not None else "",
            }
        )
    return entries


def story_kind_for_part(part_name: str) -> str | None:
    if part_name == "word/document.xml":
        return "main"
    if part_name.startswith("word/header"):
        return "header"
    if part_name.startswith("word/footer"):
        return "footer"
    if part_name == "word/comments.xml":
        return "comments"
    if part_name == "word/footnotes.xml":
        return "footnotes"
    if part_name == "word/endnotes.xml":
        return "endnotes"
    return None


def drawing_entries_from_xml_parts(xml_parts: dict[str, str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for part_name, text in sorted(xml_parts.items()):
        root = parse_xml(text)
        if root is None:
            continue
        for index, doc_pr in enumerate(root.findall(".//wp:docPr", NS), start=1):
            entries.append(
                {
                    "part_name": part_name,
                    "drawing_index_in_part": index,
                    "id": attr(doc_pr, None, "id"),
                    "name": attr(doc_pr, None, "name"),
                    "title": attr(doc_pr, None, "title"),
                    "description": attr(doc_pr, None, "descr"),
                    "alt_text_present": bool(attr(doc_pr, None, "title") or attr(doc_pr, None, "descr")),
                }
            )
    return entries


def hyperlink_entries_from_xml_parts(xml_parts: dict[str, str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for part_name, text in sorted(xml_parts.items()):
        root = parse_xml(text)
        if root is None:
            continue
        for index, hyperlink in enumerate(root.findall(".//w:hyperlink", NS), start=1):
            entries.append(
                {
                    "part_name": part_name,
                    "hyperlink_index_in_part": index,
                    "relationship_id": attr(hyperlink, "r", "id"),
                    "anchor": attr(hyperlink, "w", "anchor"),
                    "tooltip": attr(hyperlink, "w", "tooltip"),
                    "history": attr(hyperlink, "w", "history"),
                    "text": element_text(hyperlink),
                }
            )
    return entries


def style_evidence_from_xml(styles_xml: str) -> dict[str, Any]:
    root = parse_xml(styles_xml)
    if root is None:
        return {
            "style_count": 0,
            "style_type_counts": {},
            "default_style_ids": {},
            "document_default_language": None,
        }

    styles = root.findall("w:style", NS)
    type_counts: dict[str, int] = {}
    default_style_ids: dict[str, str | None] = {}
    for style in styles:
        style_type = attr(style, "w", "type") or "unknown"
        type_counts[style_type] = type_counts.get(style_type, 0) + 1
        if attr(style, "w", "default") in {"1", "true"}:
            default_style_ids[style_type] = attr(style, "w", "styleId")
    default_lang = root.find("w:docDefaults/w:rPrDefault/w:rPr/w:lang", NS)
    return {
        "style_count": len(styles),
        "style_type_counts": type_counts,
        "default_style_ids": default_style_ids,
        "document_default_language": attr(default_lang, "w", "val"),
        "theme_font_language": attr(default_lang, "w", "theme"),
    }


def numbering_evidence_from_xml(numbering_xml: str) -> dict[str, Any]:
    root = parse_xml(numbering_xml)
    if root is None:
        return {
            "abstract_numbering_count": 0,
            "numbering_definition_count": 0,
            "level_count": 0,
            "formats": {},
        }

    formats: dict[str, int] = {}
    for num_fmt in root.findall(".//w:numFmt", NS):
        value = attr(num_fmt, "w", "val") or "unknown"
        formats[value] = formats.get(value, 0) + 1
    return {
        "abstract_numbering_count": len(root.findall("w:abstractNum", NS)),
        "numbering_definition_count": len(root.findall("w:num", NS)),
        "level_count": len(root.findall(".//w:lvl", NS)),
        "formats": formats,
    }


def settings_evidence_from_xml(settings_xml: str) -> dict[str, Any]:
    root = parse_xml(settings_xml)
    if root is None:
        return {
            "settings_present": False,
            "track_revisions": None,
            "document_protection_present": None,
            "even_and_odd_headers": None,
            "update_fields": None,
        }
    return {
        "settings_present": True,
        "track_revisions": root.find("w:trackRevisions", NS) is not None,
        "document_protection_present": root.find("w:documentProtection", NS) is not None,
        "even_and_odd_headers": root.find("w:evenAndOddHeaders", NS) is not None,
        "update_fields": attr(root.find("w:updateFields", NS), "w", "val"),
        "compatibility_setting_count": len(root.findall(".//w:compatSetting", NS)),
    }


def extended_properties_from_xml(app_xml: str) -> dict[str, Any]:
    root = parse_xml(app_xml)
    if root is None:
        return {}
    values = {}
    for child in list(root):
        local = child.tag.rsplit("}", 1)[-1]
        if local in {"Pages", "Words", "Characters", "Lines", "Paragraphs", "Company", "Application", "DocSecurity"}:
            values[local[0].lower() + local[1:]] = child.text
    return values


def font_evidence_from_xml_parts(xml_parts: dict[str, str]) -> dict[str, Any]:
    font_table_entries = font_table_entries_from_xml(xml_parts.get("word/fontTable.xml", ""))
    theme_font_entries = theme_font_entries_from_xml(xml_parts.get("word/theme/theme1.xml", ""))
    style_font_entries = style_font_entries_from_xml(xml_parts.get("word/styles.xml", ""))
    direct_run_font_entries = direct_run_font_entries_from_xml_parts(xml_parts)
    font_name_counts = Counter()

    for entry in font_table_entries:
        add_counter_value(font_name_counts, entry.get("name"))
    for entry in theme_font_entries:
        add_counter_value(font_name_counts, entry.get("typeface"))
    for entry in style_font_entries:
        for value in entry.get("fonts", {}).values():
            add_counter_value(font_name_counts, value)
    for entry in direct_run_font_entries:
        for value in entry.get("fonts", {}).values():
            add_counter_value(font_name_counts, value)

    return {
        "font_table_entries": font_table_entries,
        "theme_font_entries": theme_font_entries,
        "style_font_entries": style_font_entries,
        "direct_run_font_entries": direct_run_font_entries,
        "font_family_counts": [
            {"font_name": font_name, "reference_count": count}
            for font_name, count in sorted(font_name_counts.items(), key=lambda item: (-item[1], item[0].lower()))
        ],
        "notes": [
            "DOCX fonts are often inherited through styles, theme fonts, and the font table; direct python-docx run.font.name is frequently null.",
            "Theme aliases such as majorHAnsi/minorHAnsi are preserved as references because their concrete font depends on the theme font scheme.",
        ],
    }


def font_table_entries_from_xml(font_table_xml: str) -> list[dict[str, Any]]:
    root = parse_xml(font_table_xml)
    if root is None:
        return []
    entries: list[dict[str, Any]] = []
    for font in root.findall("w:font", NS):
        family = font.find("w:family", NS)
        pitch = font.find("w:pitch", NS)
        charset = font.find("w:charset", NS)
        entries.append(
            {
                "name": attr(font, "w", "name"),
                "family": attr(family, "w", "val"),
                "pitch": attr(pitch, "w", "val"),
                "charset": attr(charset, "w", "val"),
                "panose1": attr(font.find("w:panose1", NS), "w", "val"),
                "signature": element_attributes(font.find("w:sig", NS), "w"),
            }
        )
    return entries


def theme_font_entries_from_xml(theme_xml: str) -> list[dict[str, Any]]:
    root = parse_xml(theme_xml)
    if root is None:
        return []
    entries: list[dict[str, Any]] = []
    font_scheme = root.find(".//a:fontScheme", NS)
    if font_scheme is None:
        return entries

    for scheme_kind in ("majorFont", "minorFont"):
        scheme = font_scheme.find(f"a:{scheme_kind}", NS)
        if scheme is None:
            continue
        for role in ("latin", "ea", "cs"):
            role_element = scheme.find(f"a:{role}", NS)
            if role_element is not None:
                entries.append(
                    {
                        "scheme": scheme_kind,
                        "role": role,
                        "typeface": attr(role_element, None, "typeface"),
                        "script": None,
                    }
                )
        for supplemental in scheme.findall("a:font", NS):
            entries.append(
                {
                    "scheme": scheme_kind,
                    "role": "supplemental",
                    "typeface": attr(supplemental, None, "typeface"),
                    "script": attr(supplemental, None, "script"),
                }
            )
    return entries


def style_font_entries_from_xml(styles_xml: str) -> list[dict[str, Any]]:
    root = parse_xml(styles_xml)
    if root is None:
        return []
    entries: list[dict[str, Any]] = []
    for style in root.findall("w:style", NS):
        r_fonts = style.find("w:rPr/w:rFonts", NS)
        if r_fonts is None:
            continue
        fonts = rfonts_attributes(r_fonts)
        if not fonts:
            continue
        entries.append(
            {
                "style_id": attr(style, "w", "styleId"),
                "style_type": attr(style, "w", "type"),
                "style_name": attr(style.find("w:name", NS), "w", "val"),
                "based_on": attr(style.find("w:basedOn", NS), "w", "val"),
                "fonts": fonts,
            }
        )
    return entries


def direct_run_font_entries_from_xml_parts(xml_parts: dict[str, str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for part_name, text in sorted(xml_parts.items()):
        if not part_name.endswith(".xml") or classify_part_name(part_name) in {"font_table", "theme", "styles"}:
            continue
        root = parse_xml(text)
        if root is None:
            continue
        for index, r_fonts in enumerate(root.findall(".//w:rPr/w:rFonts", NS), start=1):
            fonts = rfonts_attributes(r_fonts)
            if not fonts:
                continue
            entries.append(
                {
                    "part_name": part_name,
                    "run_font_index_in_part": index,
                    "part_kind": classify_part_name(part_name),
                    "fonts": fonts,
                }
            )
    return entries


def rfonts_attributes(r_fonts: ET.Element) -> dict[str, str]:
    keys = (
        "ascii",
        "hAnsi",
        "eastAsia",
        "cs",
        "asciiTheme",
        "hAnsiTheme",
        "eastAsiaTheme",
        "cstheme",
    )
    values: dict[str, str] = {}
    for key in keys:
        value = attr(r_fonts, "w", key)
        if value:
            values[key] = value
    return values


def element_attributes(element: ET.Element | None, prefix: str | None = None) -> dict[str, str]:
    if element is None:
        return {}
    values: dict[str, str] = {}
    namespace = NS.get(prefix) if prefix else None
    for key, value in element.attrib.items():
        if namespace and key.startswith(f"{{{namespace}}}"):
            values[key.rsplit("}", 1)[-1]] = value
        else:
            values[key] = value
    return values


def add_counter_value(counter: Counter, value: str | None) -> None:
    if value:
        counter[value] += 1


def paragraph_entries_from_document(document: Any, *, include_raw_xml: bool = False) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text or ""
        style_name = paragraph.style.name if paragraph.style is not None else None
        entries.append(
            {
                "paragraph_index": index,
                "text": text,
                "text_character_count": len(text),
                "style_name": style_name,
                "alignment": paragraph.alignment.name if paragraph.alignment is not None else None,
                "run_count": len(paragraph.runs),
                "hyperlink_count": paragraph._p.xml.count("<w:hyperlink"),
                "numbering_detected": "<w:numPr" in paragraph._p.xml,
                "page_break_before": bool(paragraph.paragraph_format.page_break_before),
                "keep_with_next": bool(paragraph.paragraph_format.keep_with_next),
                "runs": [
                    run_entry(run, run_index, include_raw_xml=include_raw_xml)
                    for run_index, run in enumerate(paragraph.runs, start=1)
                ],
            }
        )
    return entries


def run_entry(run: Any, run_index: int, *, include_raw_xml: bool = False) -> dict[str, Any]:
    font = run.font
    entry = {
        "run_index": run_index,
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline) if run.underline is not None else None,
        "font_name": font.name,
        "font_size_pt": font.size.pt if font.size is not None else None,
        "style_name": run.style.name if run.style is not None else None,
        "contains_break": "<w:br" in run._r.xml,
        "contains_drawing": "<w:drawing" in run._r.xml,
        "raw_xml_sha256": text_sha256(run._r.xml),
    }
    if include_raw_xml:
        entry["raw_xml"] = run._r.xml
    return entry


def table_entries_from_document(document: Any) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    table_entries: list[dict[str, Any]] = []
    cell_entries: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = table.rows
        column_count = max((len(row.cells) for row in rows), default=0)
        table_entries.append(
            {
                "table_index": table_index,
                "style_name": table.style.name if table.style is not None else None,
                "row_count": len(rows),
                "column_count": column_count,
                "text_sample": table_text_sample(table),
            }
        )
        for row_index, row in enumerate(rows, start=1):
            for column_index, cell in enumerate(row.cells, start=1):
                cell_entries.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "column_index": column_index,
                        "text": cell.text,
                        "paragraph_count": len(cell.paragraphs),
                        "nested_table_count": len(cell.tables),
                    }
                )
    return table_entries, cell_entries


def table_text_sample(table: Any) -> list[list[str]]:
    sample: list[list[str]] = []
    for row in table.rows[:5]:
        sample.append([cell.text for cell in row.cells[:8]])
    return sample


def section_entries_from_document(document: Any) -> list[dict[str, Any]]:
    entries = []
    for index, section in enumerate(document.sections, start=1):
        entries.append(
            {
                "section_index": index,
                "page_width_inches": length_inches(section.page_width),
                "page_height_inches": length_inches(section.page_height),
                "orientation": section.orientation.name if section.orientation is not None else None,
                "top_margin_inches": length_inches(section.top_margin),
                "bottom_margin_inches": length_inches(section.bottom_margin),
                "left_margin_inches": length_inches(section.left_margin),
                "right_margin_inches": length_inches(section.right_margin),
                "header_distance_inches": length_inches(section.header_distance),
                "footer_distance_inches": length_inches(section.footer_distance),
            }
        )
    return entries


def style_entries_from_document(document: Any) -> list[dict[str, Any]]:
    entries = []
    for style in document.styles:
        entries.append(
            {
                "style_id": style.style_id,
                "name": style.name,
                "type": style.type.name if style.type is not None else None,
                "builtin": style.builtin,
                "hidden": style.hidden,
                "quick_style": style.quick_style,
            }
        )
    return entries


def media_entries_from_package(package: dict[str, Any]) -> list[dict[str, Any]]:
    entries = []
    for part in package["part_entries"]:
        if part["part_name"].startswith("word/media/"):
            entries.append(
                {
                    "part_name": part["part_name"],
                    "file_name": Path(part["part_name"]).name,
                    "byte_size": part["byte_size"],
                    "suffix": part["suffix"],
                }
            )
    return entries


def header_footer_entries_from_package(package: dict[str, Any]) -> list[dict[str, Any]]:
    entries = []
    for metrics in package["xml_part_metrics"]:
        part_name = metrics["part_name"]
        if part_name.startswith("word/header") or part_name.startswith("word/footer"):
            entries.append(
                {
                    "part_name": part_name,
                    "part_kind": "header" if part_name.startswith("word/header") else "footer",
                    "paragraph_tag_count": metrics["paragraph_tag_count"],
                    "table_tag_count": metrics["table_tag_count"],
                    "drawing_tag_count": metrics["drawing_tag_count"],
                    "text_tag_count": metrics["text_tag_count"],
                }
            )
    return entries


def core_properties(document: Any) -> dict[str, Any]:
    props = document.core_properties
    return {
        "author": props.author,
        "category": props.category,
        "comments": props.comments,
        "content_status": props.content_status,
        "created": serialize_value(props.created),
        "identifier": props.identifier,
        "keywords": props.keywords,
        "language": props.language,
        "last_modified_by": props.last_modified_by,
        "last_printed": serialize_value(props.last_printed),
        "modified": serialize_value(props.modified),
        "revision": props.revision,
        "subject": props.subject,
        "title": props.title,
        "version": props.version,
    }


def document_warnings(document: Any, summary: dict[str, Any], style_evidence: dict[str, Any]) -> list[dict[str, Any]]:
    warnings: list[dict[str, Any]] = []
    props = document.core_properties
    if not props.title:
        warnings.append({"warning_code": "DOCX_TITLE_MISSING", "severity": "warning", "message": "DOCX core title is empty."})
    if not props.language and not style_evidence.get("document_default_language"):
        warnings.append({"warning_code": "DOCX_LANGUAGE_MISSING", "severity": "warning", "message": "DOCX core language is empty."})
    elif not props.language and style_evidence.get("document_default_language"):
        warnings.append(
            {
                "warning_code": "DOCX_CORE_LANGUAGE_EMPTY_DEFAULT_LANGUAGE_FOUND",
                "severity": "info",
                "message": "DOCX core language is empty, but WordprocessingML styles declare a default language.",
                "evidence": {"document_default_language": style_evidence.get("document_default_language")},
            }
        )
    if summary["image_count"] > 0:
        warnings.append(
            {
                "warning_code": "DOCX_IMAGES_REQUIRE_ALT_REVIEW",
                "severity": "review",
                "message": "DOCX contains images. Alt text and figure semantics need review.",
            }
        )
    if summary.get("drawing_count", 0) > summary.get("drawing_with_alt_text_count", 0):
        warnings.append(
            {
                "warning_code": "DOCX_DRAWING_ALT_TEXT_INCOMPLETE",
                "severity": "review",
                "message": "Some WordprocessingML drawings do not expose title or description attributes.",
            }
        )
    return warnings


def parse_xml(text: str) -> ET.Element | None:
    if not text:
        return None
    try:
        return ET.fromstring(text.encode("utf-8"))
    except ET.ParseError:
        return None


def qn(prefix: str, local_name: str) -> str:
    return f"{{{NS[prefix]}}}{local_name}"


def attr(element: ET.Element | None, prefix: str | None, local_name: str) -> str | None:
    if element is None:
        return None
    attr_name = local_name if prefix is None else qn(prefix, local_name)
    return element.attrib.get(attr_name)


def parse_int_attr(element: ET.Element | None, prefix: str, local_name: str) -> int | None:
    value = attr(element, prefix, local_name)
    if value is None:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def element_text(element: ET.Element | None) -> str:
    if element is None:
        return ""
    return "".join(text.text or "" for text in element.iter(qn("w", "t")))


def element_xml(element: ET.Element | None) -> str | None:
    if element is None:
        return None
    return ET.tostring(element, encoding="unicode")


def text_sha256(value: str | None) -> str | None:
    if value is None:
        return None
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def length_inches(value: Any) -> float | None:
    return round(value.inches, 4) if value is not None else None


def serialize_value(value: Any) -> Any:
    if isinstance(value, datetime | date):
        return value.isoformat()
    return value


def count_regex(text: str, pattern: str) -> int:
    return len(re.findall(pattern, text))
