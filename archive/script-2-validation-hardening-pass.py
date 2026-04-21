"""
Script 2 — Image-retaining DOCX remediation and reconstruction.

Purpose
- Consume the Script 1 companion DOCX, the authoritative JSON manifest, and slide images.
- Build the final image-retaining DOCX.
- Improve Word structure where the preserved evidence supports it.
- Preserve unresolved warnings and manual-review needs honestly.
- Produce a remediation report.

What this script does
- Uses the JSON manifest as the authoritative process truth.
- Validates that the candidate DOCX exists and that its declared pairing with the manifest is plausible.
- Resolves slide images and inserts them into the final DOCX.
- Rebuilds each slide as:
  - slide heading
  - slide image
  - slide text as real Word text
  - visual descriptions
  - manual-review notes when needed
- Reconstructs stronger Word structure where safe.

What this script does not do
- It does not perform OCR.
- It does not re-parse the PPTX as its main workflow.
- It does not certify WCAG compliance.
- It does not silently infer diagram meaning or connector meaning that was not preserved.
- It does not treat the candidate DOCX as the main semantic source.

Dependencies
- python-docx

Basic usage
- Edit the CONFIG block below.
- Set INPUT_DOCX_PATH to the Script 1 companion DOCX.
- Set INPUT_MANIFEST_JSON_PATH to the authoritative JSON manifest from Script 1.
- Optionally set INPUT_REVIEW_NOTES_PATH to the Script 1 review notes Markdown.
- Set SLIDE_IMAGE_DIR if you want to override the image directory from the manifest.
- Run the script directly. No CLI arguments are required.

Manual-review expectations
- This script improves the DOCX structure where confidence is high enough.
- It preserves warnings rather than suppressing them.
- Final review still belongs to the actual remediated DOCX artifact.
"""

from __future__ import annotations

import json
import logging
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


# ==========================
# CONFIG — CORE INPUT/OUTPUT
# ==========================
INPUT_DOCX_PATH = "/Users/computerk/Library/CloudStorage/Dropbox/0_VSCode/extracted_chat_artifacts/test_outputs/Disorders_in_the_DSM5TR_Presentation_companion.docx"
INPUT_MANIFEST_JSON_PATH = "/Users/computerk/Library/CloudStorage/Dropbox/0_VSCode/extracted_chat_artifacts/test_outputs/Disorders_in_the_DSM5TR_Presentation_manifest.json"
INPUT_REVIEW_NOTES_PATH = "/Users/computerk/Library/CloudStorage/Dropbox/0_VSCode/extracted_chat_artifacts/test_outputs/Disorders_in_the_DSM5TR_Presentation_review_notes.md"

OUTPUT_DOCX_PATH = "/Users/computerk/Library/CloudStorage/Dropbox/0_VSCode/extracted_chat_artifacts/test_outputs/Disorders_in_the_DSM5TR_Presentation_final.docx"
OUTPUT_REMEDIATION_REPORT_PATH = "/Users/computerk/Library/CloudStorage/Dropbox/0_VSCode/extracted_chat_artifacts/test_outputs/Disorders_in_the_DSM5TR_Presentation_remediation_report_final.md"

OVERWRITE = False
DRY_RUN = False

# ==========================
# CONFIG — IMAGE HANDLING
# ==========================
SLIDE_IMAGE_DIR = ""
REQUIRE_SLIDE_IMAGES = True

# ==========================
# CONFIG — VALIDATION
# ==========================
REQUIRE_FINAL_FILE_MATCH = True
REQUIRE_REVIEW_NOTES = False
FAIL_ON_UNRESOLVED_CRITICAL_WARNINGS = False

# ==========================
# CONFIG — DOCUMENT / OUTPUT
# ==========================
DOC_TITLE = ""
SLIDE_SEPARATOR_MODE = "PAGE_BREAK"  # "PAGE_BREAK" | "BLANK_LINE"
INCLUDE_WARNING_NOTES = True
INCLUDE_SOURCE_QUALITY_NOTES = True
INCLUDE_SLIDE_TEXT_SECTION = True
INCLUDE_VISUAL_DESCRIPTIONS_SECTION = True

# ==========================
# CONFIG — LOGGING
# ==========================
LOG_LEVEL = "INFO"  # "DEBUG" | "INFO" | "WARNING" | "ERROR"
LOG_TO_FILE = False
LOG_FILE_PATH = "./script2_remediation.log"
SHOW_PROGRESS = True

# Input and output rules
# - INPUT_DOCX_PATH must point to the Script 1 companion DOCX.
# - INPUT_MANIFEST_JSON_PATH must point to the authoritative JSON manifest from Script 1.
# - Script 2 uses the manifest as the process truth and uses the candidate DOCX mainly for
#   existence and declared-pairing validation, not as the main semantic reconstruction source.
# - Slide images are required by default. They are resolved from SLIDE_IMAGE_DIR first if set,
#   otherwise from the manifest preview-image data.
# - Preview-image candidates in the manifest must be relative filenames or relative paths.
# - If output paths are blank, they are derived from the input DOCX stem.
# - When DRY_RUN=True, the script validates inputs, parses the manifest, resolves image paths,
#   builds output in memory, and prints summary information, but makes zero filesystem changes.


STATUS_SUCCESS = "SUCCESS"
STATUS_SUCCESS_WITH_WARNINGS = "SUCCESS_WITH_WARNINGS"
STATUS_DRY_RUN_COMPLETE = "DRY_RUN_COMPLETE"
STATUS_FAILURE = "FAILURE"

EXIT_SUCCESS = 0
EXIT_CONFIG_ERROR = 2
EXIT_RUNTIME_ERROR = 3

ALLOWED_LOG_LEVELS = {"DEBUG", "INFO", "WARNING", "ERROR"}
ALLOWED_SLIDE_SEPARATOR_MODES = {"PAGE_BREAK", "BLANK_LINE"}

SOURCE_TEXT_CONTAINER_TYPES = {"placeholder_text", "standalone_text_box", "table_cell"}
INTERPRETED_TEXT_ROLES = {
    "title_candidate",
    "body_content",
    "caption",
    "callout",
    "label",
    "annotation",
    "unknown",
}
ROLE_CONFIDENCE_VALUES = {"high", "medium", "low"}
LIST_KIND_VALUES = {"plain", "bullet", "number"}
DESCRIPTION_SOURCE_VALUES = {"descr", "title", "missing"}
PROJECTION_STATUS_VALUES = {"planned", "written", "omitted", "manual_only"}
SUPPORTED_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"}

BASE_CRITICAL_WARNING_CODES = {
    "CONNECTOR_RELATIONSHIP_WARNING",
    "GROUPED_CONTENT_DETECTED",
    "READING_ORDER_LOW_CONFIDENCE",
    "DESCRIPTION_METADATA_MISSING",
    "NO_EXTRACTABLE_TEXT",
}

IMAGE_CRITICAL_WARNING_CODES = {
    "SLIDE_IMAGE_DIRECTORY_UNAVAILABLE",
    "SLIDE_IMAGE_NOT_FOUND",
    "SLIDE_IMAGE_INSERTION_FAILED",
}

LOGGER = logging.getLogger("script2_remediation")
WCAG_AAA_TEXT_COLOR = RGBColor(0, 0, 0)
WCAG_AAA_TEXT_COLOR_HEX = "000000"


class ConfigError(Exception):
    """Raised when CONFIG values are invalid or incompatible."""


class ManifestError(Exception):
    """Raised when the manifest is missing required workflow content."""


@dataclass
class OutputPaths:
    docx: Path
    report_md: Path


@dataclass
class ImageResolutionResult:
    slide_number: int
    image_path: Optional[Path]
    warnings: list[dict[str, Any]]


def critical_warning_codes() -> set[str]:
    if REQUIRE_SLIDE_IMAGES:
        return BASE_CRITICAL_WARNING_CODES | IMAGE_CRITICAL_WARNING_CODES
    return set(BASE_CRITICAL_WARNING_CODES)


def setup_logging() -> logging.Logger:
    logger = LOGGER
    logger.handlers.clear()
    logger.setLevel(getattr(logging, LOG_LEVEL, logging.INFO))
    logger.propagate = False

    formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")

    stream_handler = logging.StreamHandler(sys.stderr)
    stream_handler.setFormatter(formatter)
    logger.addHandler(stream_handler)

    if LOG_TO_FILE and not DRY_RUN:
        log_path = Path(LOG_FILE_PATH)
        log_path.parent.mkdir(parents=True, exist_ok=True)
        file_handler = logging.FileHandler(log_path, encoding="utf-8")
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

    return logger


def fail(message: str, exit_code: int) -> int:
    LOGGER.error(message)
    return exit_code


def make_warning(code: str, message: str, scope: str, manual_review_required: bool) -> dict[str, Any]:
    return {
        "warning_code": code,
        "warning_message": message,
        "warning_scope": scope,
        "manual_review_required": manual_review_required,
    }


def dedupe_warning_entries(entries: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[tuple[str, str, str]] = set()
    output: list[dict[str, Any]] = []
    for entry in entries:
        key = (entry["warning_code"], entry["warning_message"], entry["warning_scope"])
        if key in seen:
            continue
        seen.add(key)
        output.append(entry)
    return output


def derive_output_paths(input_docx_path: Path) -> OutputPaths:
    stem = input_docx_path.stem
    parent = input_docx_path.parent

    output_docx = Path(OUTPUT_DOCX_PATH).expanduser() if OUTPUT_DOCX_PATH else parent / f"{stem}_remediated.docx"
    output_report = (
        Path(OUTPUT_REMEDIATION_REPORT_PATH).expanduser()
        if OUTPUT_REMEDIATION_REPORT_PATH
        else parent / f"{stem}_remediation_report.md"
    )
    return OutputPaths(docx=output_docx, report_md=output_report)


def validate_config() -> tuple[Path, Path, Optional[Path], OutputPaths]:
    input_docx = Path(INPUT_DOCX_PATH).expanduser()
    input_manifest = Path(INPUT_MANIFEST_JSON_PATH).expanduser()
    input_review_notes = Path(INPUT_REVIEW_NOTES_PATH).expanduser() if INPUT_REVIEW_NOTES_PATH else None

    if not INPUT_DOCX_PATH:
        raise ConfigError("INPUT_DOCX_PATH is empty. Set it to the Script 1 companion DOCX.")
    if not input_docx.exists() or not input_docx.is_file() or input_docx.suffix.lower() != ".docx":
        raise ConfigError("INPUT_DOCX_PATH must point to an existing .docx file.")

    if not INPUT_MANIFEST_JSON_PATH:
        raise ConfigError("INPUT_MANIFEST_JSON_PATH is empty. Set it to the Script 1 JSON manifest.")
    if not input_manifest.exists() or not input_manifest.is_file() or input_manifest.suffix.lower() != ".json":
        raise ConfigError("INPUT_MANIFEST_JSON_PATH must point to an existing .json file.")

    if REQUIRE_REVIEW_NOTES and input_review_notes is None:
        raise ConfigError("REQUIRE_REVIEW_NOTES is True but INPUT_REVIEW_NOTES_PATH is empty.")
    if input_review_notes is not None and (not input_review_notes.exists() or not input_review_notes.is_file()):
        raise ConfigError("INPUT_REVIEW_NOTES_PATH must point to an existing file when provided.")

    if LOG_LEVEL not in ALLOWED_LOG_LEVELS:
        allowed = ", ".join(sorted(ALLOWED_LOG_LEVELS))
        raise ConfigError(f"LOG_LEVEL must be one of: {allowed}.")
    if SLIDE_SEPARATOR_MODE not in ALLOWED_SLIDE_SEPARATOR_MODES:
        allowed = ", ".join(sorted(ALLOWED_SLIDE_SEPARATOR_MODES))
        raise ConfigError(f"SLIDE_SEPARATOR_MODE must be one of: {allowed}.")
    if LOG_TO_FILE and not LOG_FILE_PATH:
        raise ConfigError("LOG_TO_FILE is True but LOG_FILE_PATH is empty.")

    output_paths = derive_output_paths(input_docx)
    for path in [output_paths.docx, output_paths.report_md]:
        if path.suffix == "":
            raise ConfigError(f"Output path must be a file path, not a directory: {path}")
        if path.exists() and path.is_dir():
            raise ConfigError(f"Output path points to an existing directory: {path}")
        if path.parent.exists() and not path.parent.is_dir():
            raise ConfigError(f"Output path parent exists but is not a directory: {path.parent}")
        if path.exists() and not OVERWRITE and not DRY_RUN:
            raise ConfigError(f"Output file already exists. Set OVERWRITE = True or change the path: {path}")

    return input_docx, input_manifest, input_review_notes, output_paths


def ensure_output_directories(output_paths: OutputPaths) -> None:
    for path in [output_paths.docx, output_paths.report_md]:
        if path.parent.exists():
            continue
        if DRY_RUN:
            LOGGER.info("DRY_RUN: would create directory %s", path.parent)
        else:
            path.parent.mkdir(parents=True, exist_ok=True)


def load_manifest(manifest_path: Path) -> dict[str, Any]:
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise ManifestError(f"Failed to read JSON manifest '{manifest_path}': {exc}") from exc

    validate_manifest_structure(manifest)
    return manifest


def require_key(obj: dict[str, Any], key: str, context: str) -> Any:
    if key not in obj:
        raise ManifestError(f"Manifest is missing required key '{key}' in {context}.")
    return obj[key]


def require_type(value: Any, expected_type: type | tuple[type, ...], context: str) -> None:
    if not isinstance(value, expected_type):
        raise ManifestError(f"Manifest value at {context} has invalid type.")


def require_enum(value: Any, allowed: set[str], context: str) -> None:
    if value not in allowed:
        allowed_text = ", ".join(sorted(allowed))
        raise ManifestError(f"Manifest value at {context} must be one of: {allowed_text}.")


def validate_manifest_structure(manifest: dict[str, Any]) -> None:
    top_required = {
        "manifest_version",
        "manifest_kind",
        "source_package",
        "target_package",
        "manifest_summary",
        "slide_entries",
    }
    missing = sorted(top_required - set(manifest.keys()))
    if missing:
        raise ManifestError(f"Manifest is missing required top-level keys: {', '.join(missing)}")

    if manifest["manifest_kind"] != "pptx_to_docx_remediation_manifest":
        raise ManifestError("Manifest kind is not 'pptx_to_docx_remediation_manifest'.")

    source_package = manifest["source_package"]
    target_package = manifest["target_package"]
    require_type(source_package, dict, "source_package")
    require_type(target_package, dict, "target_package")

    for context_name, pkg in [("source_package", source_package), ("target_package", target_package)]:
        require_type(require_key(pkg, "format_family", context_name), str, f"{context_name}.format_family")
        require_type(require_key(pkg, "package_model", context_name), str, f"{context_name}.package_model")
        require_type(require_key(pkg, "document_vocabulary", context_name), str, f"{context_name}.document_vocabulary")

    if "run_config" in manifest:
        require_type(manifest["run_config"], dict, "run_config")

    summary = manifest["manifest_summary"]
    require_type(summary, dict, "manifest_summary")

    slide_entries = manifest["slide_entries"]
    require_type(slide_entries, list, "slide_entries")

    seen_slide_numbers: set[int] = set()
    for slide_index, slide in enumerate(slide_entries, start=1):
        slide_number = validate_slide_entry(slide, f"slide_entries[{slide_index}]")
        if slide_number in seen_slide_numbers:
            raise ManifestError(f"Manifest contains duplicate slide_number value: {slide_number}")
        seen_slide_numbers.add(slide_number)


def validate_slide_entry(slide: dict[str, Any], context: str) -> int:
    require_type(slide, dict, context)
    slide_number = require_key(slide, "slide_number", context)
    require_type(slide_number, int, f"{context}.slide_number")
    if slide_number < 1:
        raise ManifestError(f"Manifest value at {context}.slide_number must be >= 1.")

    observed_source = require_key(slide, "observed_source", context)
    normalized_workflow = require_key(slide, "normalized_workflow", context)
    projected_target = require_key(slide, "projected_target", context)
    warning_entries = require_key(slide, "warning_entries", context)
    text_block_entries = require_key(slide, "text_block_entries", context)
    visual_entries = require_key(slide, "visual_entries", context)

    require_type(observed_source, dict, f"{context}.observed_source")
    require_type(normalized_workflow, dict, f"{context}.normalized_workflow")
    require_type(projected_target, dict, f"{context}.projected_target")
    require_type(warning_entries, list, f"{context}.warning_entries")
    require_type(text_block_entries, list, f"{context}.text_block_entries")
    require_type(visual_entries, list, f"{context}.visual_entries")

    require_type(require_key(normalized_workflow, "chosen_title", f"{context}.normalized_workflow"), str, f"{context}.normalized_workflow.chosen_title")
    require_enum(
        require_key(normalized_workflow, "title_source", f"{context}.normalized_workflow"),
        {"TITLE_PLACEHOLDER", "FIRST_TEXT_LINE", "UNTITLED_FALLBACK"},
        f"{context}.normalized_workflow.title_source",
    )

    require_type(
        require_key(projected_target, "projected_slide_heading_text", f"{context}.projected_target"),
        str,
        f"{context}.projected_target.projected_slide_heading_text",
    )
    require_type(
        require_key(projected_target, "projected_slide_heading_style", f"{context}.projected_target"),
        str,
        f"{context}.projected_target.projected_slide_heading_style",
    )
    require_type(
        require_key(projected_target, "projected_slide_separator_mode", f"{context}.projected_target"),
        str,
        f"{context}.projected_target.projected_slide_separator_mode",
    )
    require_enum(
        require_key(projected_target, "projection_status", f"{context}.projected_target"),
        PROJECTION_STATUS_VALUES,
        f"{context}.projected_target.projection_status",
    )

    preview_candidates = observed_source.get("preview_image_candidates", [])
    require_type(preview_candidates, list, f"{context}.observed_source.preview_image_candidates")
    for candidate_index, candidate in enumerate(preview_candidates, start=1):
        require_type(candidate, str, f"{context}.observed_source.preview_image_candidates[{candidate_index}]")
        validate_preview_candidate(candidate, f"{context}.observed_source.preview_image_candidates[{candidate_index}]")

    for warning_index, warning in enumerate(warning_entries, start=1):
        validate_warning_entry(warning, f"{context}.warning_entries[{warning_index}]")

    for block_index, block in enumerate(text_block_entries, start=1):
        validate_text_block_entry(block, f"{context}.text_block_entries[{block_index}]")

    for visual_index, visual in enumerate(visual_entries, start=1):
        validate_visual_entry(visual, f"{context}.visual_entries[{visual_index}]")

    return slide_number


def validate_preview_candidate(candidate: str, context: str) -> None:
    candidate_path = Path(candidate)
    if candidate_path.is_absolute():
        raise ManifestError(f"Manifest value at {context} must be a relative filename or relative path, not an absolute path.")
    if candidate_path.parts and candidate_path.parts[0] == "..":
        raise ManifestError(f"Manifest value at {context} must not traverse outside the slide image directory.")


def validate_warning_entry(warning: dict[str, Any], context: str) -> None:
    require_type(warning, dict, context)
    require_type(require_key(warning, "warning_code", context), str, f"{context}.warning_code")
    require_type(require_key(warning, "warning_message", context), str, f"{context}.warning_message")
    require_type(require_key(warning, "warning_scope", context), str, f"{context}.warning_scope")
    require_type(require_key(warning, "manual_review_required", context), bool, f"{context}.manual_review_required")


def validate_text_block_entry(block: dict[str, Any], context: str) -> None:
    require_type(block, dict, context)

    observed_source = require_key(block, "observed_source", context)
    normalized_workflow = require_key(block, "normalized_workflow", context)
    projected_target = require_key(block, "projected_target", context)
    warning_entries = require_key(block, "warning_entries", context)
    paragraph_entries = require_key(block, "paragraph_entries", context)

    require_type(observed_source, dict, f"{context}.observed_source")
    require_type(normalized_workflow, dict, f"{context}.normalized_workflow")
    require_type(projected_target, dict, f"{context}.projected_target")
    require_type(warning_entries, list, f"{context}.warning_entries")
    require_type(paragraph_entries, list, f"{context}.paragraph_entries")

    source_type = require_key(observed_source, "source_text_container_type", f"{context}.observed_source")
    require_enum(source_type, SOURCE_TEXT_CONTAINER_TYPES, f"{context}.observed_source.source_text_container_type")

    role = require_key(normalized_workflow, "interpreted_text_role", f"{context}.normalized_workflow")
    require_enum(role, INTERPRETED_TEXT_ROLES, f"{context}.normalized_workflow.interpreted_text_role")

    confidence = require_key(normalized_workflow, "role_confidence", f"{context}.normalized_workflow")
    require_enum(confidence, ROLE_CONFIDENCE_VALUES, f"{context}.normalized_workflow.role_confidence")

    require_type(
        require_key(normalized_workflow, "auto_merge_allowed", f"{context}.normalized_workflow"),
        bool,
        f"{context}.normalized_workflow.auto_merge_allowed",
    )
    require_type(
        require_key(normalized_workflow, "ordered_position", f"{context}.normalized_workflow"),
        int,
        f"{context}.normalized_workflow.ordered_position",
    )

    for warning_index, warning in enumerate(warning_entries, start=1):
        validate_warning_entry(warning, f"{context}.warning_entries[{warning_index}]")

    for paragraph_index, paragraph in enumerate(paragraph_entries, start=1):
        validate_paragraph_entry(paragraph, f"{context}.paragraph_entries[{paragraph_index}]")


def validate_paragraph_entry(paragraph: dict[str, Any], context: str) -> None:
    require_type(paragraph, dict, context)

    observed_source = require_key(paragraph, "observed_source", context)
    normalized_workflow = require_key(paragraph, "normalized_workflow", context)
    projected_target = require_key(paragraph, "projected_target", context)
    warning_entries = require_key(paragraph, "warning_entries", context)

    require_type(observed_source, dict, f"{context}.observed_source")
    require_type(normalized_workflow, dict, f"{context}.normalized_workflow")
    require_type(projected_target, dict, f"{context}.projected_target")
    require_type(warning_entries, list, f"{context}.warning_entries")

    require_type(require_key(normalized_workflow, "text", f"{context}.normalized_workflow"), str, f"{context}.normalized_workflow.text")
    require_type(require_key(normalized_workflow, "level", f"{context}.normalized_workflow"), int, f"{context}.normalized_workflow.level")

    list_kind = require_key(normalized_workflow, "list_kind", f"{context}.normalized_workflow")
    require_enum(list_kind, LIST_KIND_VALUES, f"{context}.normalized_workflow.list_kind")

    require_type(
        require_key(normalized_workflow, "ordered_sequence_reconstructable", f"{context}.normalized_workflow"),
        bool,
        f"{context}.normalized_workflow.ordered_sequence_reconstructable",
    )

    projection_status = require_key(projected_target, "projection_status", f"{context}.projected_target")
    require_enum(projection_status, PROJECTION_STATUS_VALUES, f"{context}.projected_target.projection_status")

    for warning_index, warning in enumerate(warning_entries, start=1):
        validate_warning_entry(warning, f"{context}.warning_entries[{warning_index}]")


def validate_visual_entry(visual: dict[str, Any], context: str) -> None:
    require_type(visual, dict, context)

    observed_source = require_key(visual, "observed_source", context)
    normalized_workflow = require_key(visual, "normalized_workflow", context)
    projected_target = require_key(visual, "projected_target", context)
    warning_entries = require_key(visual, "warning_entries", context)

    require_type(observed_source, dict, f"{context}.observed_source")
    require_type(normalized_workflow, dict, f"{context}.normalized_workflow")
    require_type(projected_target, dict, f"{context}.projected_target")
    require_type(warning_entries, list, f"{context}.warning_entries")

    require_type(require_key(normalized_workflow, "label", f"{context}.normalized_workflow"), str, f"{context}.normalized_workflow.label")
    require_type(require_key(normalized_workflow, "visual_type", f"{context}.normalized_workflow"), str, f"{context}.normalized_workflow.visual_type")
    require_type(require_key(normalized_workflow, "description", f"{context}.normalized_workflow"), str, f"{context}.normalized_workflow.description")

    description_source = require_key(normalized_workflow, "description_source", f"{context}.normalized_workflow")
    require_enum(description_source, DESCRIPTION_SOURCE_VALUES, f"{context}.normalized_workflow.description_source")

    require_enum(
        require_key(projected_target, "projection_status", f"{context}.projected_target"),
        PROJECTION_STATUS_VALUES,
        f"{context}.projected_target.projection_status",
    )

    for warning_index, warning in enumerate(warning_entries, start=1):
        validate_warning_entry(warning, f"{context}.warning_entries[{warning_index}]")


def validate_manifest_against_inputs(manifest: dict[str, Any], input_docx: Path) -> None:
    target_package = manifest.get("target_package", {})
    draft_output_file_name = target_package.get("draft_output_file_name", "")
    if REQUIRE_FINAL_FILE_MATCH and draft_output_file_name and draft_output_file_name != input_docx.name:
        raise ManifestError(
            "Manifest target_package.draft_output_file_name does not match INPUT_DOCX_PATH. "
            f"Manifest: {draft_output_file_name} | Input: {input_docx.name}"
        )


def resolve_slide_image_directory(manifest: dict[str, Any]) -> Optional[Path]:
    if SLIDE_IMAGE_DIR:
        return Path(SLIDE_IMAGE_DIR).expanduser()

    preview_dir = manifest.get("run_config", {}).get("preview_image_dir", None)
    if preview_dir:
        return Path(preview_dir).expanduser()

    return None


def resolve_slide_image(slide: dict[str, Any], slide_image_dir: Optional[Path]) -> ImageResolutionResult:
    warnings: list[dict[str, Any]] = []
    slide_number = slide["slide_number"]
    candidates = slide.get("observed_source", {}).get("preview_image_candidates", [])

    if slide_image_dir is None:
        if REQUIRE_SLIDE_IMAGES:
            warnings.append(
                make_warning(
                    "SLIDE_IMAGE_DIRECTORY_UNAVAILABLE",
                    "Slide image directory is unavailable, so slide image resolution failed.",
                    "slide",
                    True,
                )
            )
        return ImageResolutionResult(slide_number=slide_number, image_path=None, warnings=warnings)

    for candidate in image_candidate_names(slide_number, candidates):
        candidate_path = slide_image_dir / candidate
        if candidate_path.exists() and candidate_path.is_file() and candidate_path.suffix.lower() in SUPPORTED_IMAGE_SUFFIXES:
            return ImageResolutionResult(slide_number=slide_number, image_path=candidate_path, warnings=warnings)

    if REQUIRE_SLIDE_IMAGES:
        warnings.append(
            make_warning(
                "SLIDE_IMAGE_NOT_FOUND",
                "No slide image could be resolved for this slide.",
                "slide",
                True,
            )
        )
    return ImageResolutionResult(slide_number=slide_number, image_path=None, warnings=warnings)


def image_candidate_names(slide_number: int, manifest_candidates: list[str]) -> list[str]:
    candidates: list[str] = []
    candidates.extend(manifest_candidates)

    for suffix in sorted(SUPPORTED_IMAGE_SUFFIXES):
        candidates.extend(
            [
                f"slide_{slide_number:03d}{suffix}",
                f"slide_{slide_number:02d}{suffix}",
                f"slide_{slide_number}{suffix}",
                f"slide-{slide_number:03d}{suffix}",
                f"slide-{slide_number:02d}{suffix}",
                f"slide-{slide_number}{suffix}",
                f"Slide {slide_number}{suffix}",
                f"Slide{slide_number}{suffix}",
            ]
        )

    deduped: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        if candidate in seen:
            continue
        seen.add(candidate)
        deduped.append(candidate)
    return deduped


def read_review_notes(review_notes_path: Optional[Path]) -> dict[str, Any]:
    if review_notes_path is None:
        return {
            "present": False,
            "path": None,
            "content": "",
            "warning_count": 0,
        }

    content = review_notes_path.read_text(encoding="utf-8")
    warning_count = sum(1 for line in content.splitlines() if line.strip().startswith("- "))
    return {
        "present": True,
        "path": str(review_notes_path),
        "content": content,
        "warning_count": warning_count,
    }


def create_document_from_manifest(manifest: dict[str, Any], image_results: dict[int, ImageResolutionResult]) -> Document:
    title = resolve_document_title(manifest)
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)

    slide_entries = manifest["slide_entries"]
    for index, slide in enumerate(slide_entries, start=1):
        if SHOW_PROGRESS:
            LOGGER.info("Rebuilding slide %s", slide["slide_number"])

        add_slide_heading(document, slide)
        add_slide_image(document, slide, image_results[slide["slide_number"]])
        if INCLUDE_SLIDE_TEXT_SECTION:
            add_slide_text_section(document, slide)
        if INCLUDE_VISUAL_DESCRIPTIONS_SECTION:
            add_visual_descriptions_section(document, slide)
        if INCLUDE_WARNING_NOTES:
            add_slide_warning_section(document, slide, image_results[slide["slide_number"]])

        if index < len(slide_entries):
            add_slide_separator(document)

    enforce_wcag_aaa_text_colors(document)
    return document


def resolve_document_title(manifest: dict[str, Any]) -> str:
    if DOC_TITLE.strip():
        return DOC_TITLE.strip()

    source_name = manifest.get("source_package", {}).get("input_file_name", "")
    stem = Path(source_name).stem if source_name else "document"
    stem = stem.replace("_", " ").replace("-", " ").strip()
    if stem:
        return f"{stem} remediated companion document"
    return "Remediated companion document"


def add_slide_heading(document: Document, slide: dict[str, Any]) -> None:
    heading_text = slide["projected_target"]["projected_slide_heading_text"]
    document.add_heading(heading_text, level=1)


def slide_image_alt_metadata(slide: dict[str, Any]) -> tuple[str, str]:
    slide_number = slide["slide_number"]
    chosen_title = slide["normalized_workflow"]["chosen_title"]
    title = f"Slide {slide_number} image: {chosen_title}"
    description = (
        f"Image of slide {slide_number}, titled '{chosen_title}'. "
        "The slide's extracted text and visual descriptions are provided as real Word content following this image."
    )
    return title, description


def apply_picture_alt_text(inline_shape: Any, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def enforce_wcag_aaa_text_colors(document: Document) -> None:
    for style in document.styles:
        try:
            style.font.color.rgb = WCAG_AAA_TEXT_COLOR
        except Exception:
            pass

        try:
            color_elements = style.element.xpath(".//w:rPr/w:color")
        except Exception:
            color_elements = []
        for color_element in color_elements:
            color_element.set(qn("w:val"), WCAG_AAA_TEXT_COLOR_HEX)
            for attr_name in ("w:themeColor", "w:themeTint", "w:themeShade"):
                color_element.attrib.pop(qn(attr_name), None)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = WCAG_AAA_TEXT_COLOR


def add_slide_image(document: Document, slide: dict[str, Any], image_result: ImageResolutionResult) -> None:
    if image_result.image_path is None:
        document.add_paragraph("[Slide image unavailable. Review required.]")
        return

    try:
        inline_shape = document.add_picture(str(image_result.image_path), width=Inches(6.5))
        alt_title, alt_description = slide_image_alt_metadata(slide)
        apply_picture_alt_text(inline_shape, alt_title, alt_description)
    except Exception as exc:
        slide.setdefault("warning_entries", []).append(
            make_warning(
                "SLIDE_IMAGE_INSERTION_FAILED",
                f"Slide image insertion failed: {exc}",
                "slide",
                True,
            )
        )
        document.add_paragraph("[Slide image insertion failed. Review required.]")


def add_slide_text_section(document: Document, slide: dict[str, Any]) -> None:
    document.add_heading("Slide text", level=2)

    paragraph_entries = collect_projected_paragraphs(slide["text_block_entries"])
    if not paragraph_entries:
        document.add_paragraph("No extractable visible slide text found.")
        return

    for paragraph_entry in paragraph_entries:
        rendered = render_paragraph_for_docx(paragraph_entry)
        if rendered["kind"] == "title_like":
            document.add_paragraph(rendered["text"])
        elif rendered["kind"] == "bullet":
            style_name = bullet_style_for_level(rendered["level"])
            document.add_paragraph(rendered["text"], style=style_name)
        elif rendered["kind"] == "ordered":
            style_name = number_style_for_level(rendered["level"])
            document.add_paragraph(rendered["text"], style=style_name)
        else:
            document.add_paragraph(rendered["text"])


def collect_projected_paragraphs(text_block_entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blocks = sorted(
        text_block_entries,
        key=lambda entry: entry["normalized_workflow"]["ordered_position"],
    )
    output: list[dict[str, Any]] = []
    for block in blocks:
        normalized = block["normalized_workflow"]
        source_type = block["observed_source"]["source_text_container_type"]
        role = normalized["interpreted_text_role"]
        auto_merge_allowed = normalized["auto_merge_allowed"]

        for paragraph in block["paragraph_entries"]:
            output.append(
                {
                    "source_text_container_type": source_type,
                    "interpreted_text_role": role,
                    "role_confidence": normalized["role_confidence"],
                    "auto_merge_allowed": auto_merge_allowed,
                    "paragraph": paragraph,
                    "block_warning_entries": block["warning_entries"],
                }
            )
    return output


def render_paragraph_for_docx(paragraph_container: dict[str, Any]) -> dict[str, Any]:
    paragraph = paragraph_container["paragraph"]
    normalized = paragraph["normalized_workflow"]
    text = normalized["text"]
    level = normalized["level"]
    list_kind = normalized["list_kind"]
    ordered_sequence_reconstructable = normalized["ordered_sequence_reconstructable"]

    source_text_container_type = paragraph_container["source_text_container_type"]
    role = paragraph_container["interpreted_text_role"]

    if role == "title_candidate":
        return {"kind": "title_like", "text": text, "level": level}

    if list_kind == "bullet":
        if source_text_container_type == "standalone_text_box" and role == "unknown":
            return {"kind": "plain", "text": f"• {text}", "level": level}
        return {"kind": "bullet", "text": text, "level": level}

    if list_kind == "number":
        if ordered_sequence_reconstructable:
            return {"kind": "ordered", "text": text, "level": level}
        return {"kind": "plain", "text": f"[numbered] {text}", "level": level}

    return {"kind": "plain", "text": text, "level": level}


def bullet_style_for_level(level: int) -> str:
    if level <= 0:
        return "List Bullet"
    if level == 1:
        return "List Bullet 2"
    return "List Bullet 3"


def number_style_for_level(level: int) -> str:
    if level <= 0:
        return "List Number"
    if level == 1:
        return "List Number 2"
    return "List Number 3"


def add_visual_descriptions_section(document: Document, slide: dict[str, Any]) -> None:
    visual_entries = slide.get("visual_entries", [])
    if not visual_entries:
        return

    document.add_heading("Visual descriptions", level=2)

    for visual in visual_entries:
        normalized = visual["normalized_workflow"]
        label = normalized["label"]
        visual_type = normalized["visual_type"].capitalize()
        description = normalized["description"]
        description_source = normalized["description_source"]

        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}. {visual_type}: ").bold = True

        if description:
            text = description
            if INCLUDE_SOURCE_QUALITY_NOTES and description_source == "title":
                text += " [metadata source: title]"
        else:
            text = "Description metadata missing. Review needed."

        paragraph.add_run(text)


def add_slide_warning_section(document: Document, slide: dict[str, Any], image_result: ImageResolutionResult) -> None:
    warnings = gather_slide_warnings(slide)
    warnings.extend(image_result.warnings)
    warnings = dedupe_warning_entries(warnings)

    if not warnings:
        return

    document.add_heading("Manual review notes", level=2)
    for warning in warnings:
        document.add_paragraph(
            f"{warning['warning_code']}: {warning['warning_message']}",
            style="List Bullet",
        )


def gather_slide_warnings(slide: dict[str, Any]) -> list[dict[str, Any]]:
    warnings: list[dict[str, Any]] = []
    warnings.extend(slide.get("warning_entries", []))

    for block in slide.get("text_block_entries", []):
        warnings.extend(block.get("warning_entries", []))
        for paragraph in block.get("paragraph_entries", []):
            warnings.extend(paragraph.get("warning_entries", []))

    for visual in slide.get("visual_entries", []):
        warnings.extend(visual.get("warning_entries", []))

    return dedupe_warning_entries(warnings)


def add_slide_separator(document: Document) -> None:
    if SLIDE_SEPARATOR_MODE == "PAGE_BREAK":
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
    else:
        document.add_paragraph("")


def build_remediation_report(
    manifest: dict[str, Any],
    review_notes_info: dict[str, Any],
    input_docx: Path,
    input_manifest: Path,
    output_docx_path: Path,
    image_results: dict[int, ImageResolutionResult],
) -> str:
    slide_entries = manifest["slide_entries"]
    total_warnings = sum(
        len(dedupe_warning_entries(gather_slide_warnings(slide) + image_results[slide["slide_number"]].warnings))
        for slide in slide_entries
    )

    lines = [
        "# Remediation Report",
        "",
        "## Summary",
        "",
        f"- Runtime input DOCX: {input_docx}",
        f"- Runtime input manifest: {input_manifest}",
        f"- Remediated DOCX target: {output_docx_path}",
        f"- Source PPTX from manifest: {manifest.get('source_package', {}).get('input_file_name', '')}",
        f"- Total slides: {len(slide_entries)}",
        f"- Slides with warnings: {sum(1 for slide in slide_entries if gather_slide_warnings(slide) or image_results[slide['slide_number']].warnings)}",
        f"- Total carried-forward warnings: {total_warnings}",
        f"- Script 1 review notes present: {'Yes' if review_notes_info['present'] else 'No'}",
        "",
        "## What this script improved automatically",
        "",
        "- Rebuilt the final DOCX from the authoritative manifest rather than guessing from the draft DOCX alone.",
        "- Preserved the slide image in the final DOCX for each slide where an image could be resolved.",
        "- Re-applied heading structure from preserved title information.",
        "- Reconstructed list structure where confidence was high enough.",
        "- Preserved fallback handling where ordered sequence or role interpretation remained uncertain.",
        "- Carried unresolved warnings forward into the remediated artifact and this report.",
        "",
        "## Remaining manual review needs",
        "",
        "- Review all connector and relationship warnings.",
        "- Review all standalone text box role uncertainties.",
        "- Review visuals with missing or weak metadata.",
        "- Review any slide where reading order confidence is low.",
        "- Review any slide where the slide image could not be resolved or inserted.",
        "",
    ]

    for slide in slide_entries:
        slide_number = slide["slide_number"]
        warnings = dedupe_warning_entries(gather_slide_warnings(slide) + image_results[slide_number].warnings)
        if not warnings:
            continue
        lines.append(f"## Slide {slide_number}. {slide['normalized_workflow']['chosen_title']}")
        lines.append("")
        for warning in warnings:
            lines.append(f"- {warning['warning_code']} — {warning['warning_message']}")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def atomic_write_text(path: Path, content: str) -> None:
    temp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=path.suffix or ".tmp",
            dir=str(path.parent),
            prefix=f"{path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}_",
        ) as temp_file:
            temp_path = Path(temp_file.name)

        temp_path.write_text(content, encoding="utf-8")
        temp_path.replace(path)
    except Exception as exc:
        if temp_path is not None and temp_path.exists():
            temp_path.unlink(missing_ok=True)
        raise RuntimeError(f"Failed to write file '{path}': {exc}") from exc


def atomic_write_docx(document: Document, path: Path) -> None:
    temp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".docx",
            dir=str(path.parent),
            prefix=f"{path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}_",
        ) as temp_file:
            temp_path = Path(temp_file.name)

        document.save(str(temp_path))
        temp_path.replace(path)
    except Exception as exc:
        if temp_path is not None and temp_path.exists():
            temp_path.unlink(missing_ok=True)
        raise RuntimeError(f"Failed to write DOCX '{path}': {exc}") from exc


def write_outputs(output_paths: OutputPaths, document: Document, report_text: str) -> tuple[bool, bool]:
    if DRY_RUN:
        LOGGER.info("DRY_RUN: would write remediated DOCX to %s", output_paths.docx)
        LOGGER.info("DRY_RUN: would write remediation report to %s", output_paths.report_md)
        return False, False

    atomic_write_docx(document, output_paths.docx)
    atomic_write_text(output_paths.report_md, report_text)
    return True, True


def unresolved_critical_warnings(
    manifest: dict[str, Any],
    image_results: dict[int, ImageResolutionResult],
) -> list[dict[str, Any]]:
    critical_codes = critical_warning_codes()
    critical: list[dict[str, Any]] = []
    for slide in manifest.get("slide_entries", []):
        slide_number = slide["slide_number"]
        for warning in gather_slide_warnings(slide) + image_results[slide_number].warnings:
            if warning["warning_code"] in critical_codes:
                critical.append(warning)
    return dedupe_warning_entries(critical)


def build_run_summary(
    status: str,
    message: str,
    input_docx: Path,
    manifest_path: Path,
    review_notes_path: Optional[Path],
    output_paths: OutputPaths,
    manifest: dict[str, Any],
    outputs_written: tuple[bool, bool],
    image_results: dict[int, ImageResolutionResult],
) -> dict[str, Any]:
    docx_written, report_written = outputs_written
    return {
        "status": status,
        "message": message,
        "input_docx": str(input_docx),
        "input_manifest": str(manifest_path),
        "input_review_notes": str(review_notes_path) if review_notes_path else "",
        "output_docx": str(output_paths.docx),
        "output_report": str(output_paths.report_md),
        "docx_written": docx_written,
        "report_written": report_written,
        "slides_processed": len(manifest.get("slide_entries", [])),
        "slides_with_warnings": sum(
            1 for slide in manifest.get("slide_entries", []) if gather_slide_warnings(slide) or image_results[slide["slide_number"]].warnings
        ),
        "detected_visual_entries": manifest.get("manifest_summary", {}).get("detected_visual_entries", 0),
        "critical_warnings": len(unresolved_critical_warnings(manifest, image_results)),
        "slides_missing_images": sum(1 for result in image_results.values() if result.image_path is None),
    }


def print_run_summary(summary: dict[str, Any]) -> None:
    lines = [
        f"Run status: {summary['status']}",
        f"Message: {summary['message']}",
        f"Dry run: {DRY_RUN}",
        f"Input DOCX: {summary['input_docx']}",
        f"Input manifest: {summary['input_manifest']}",
        f"Input review notes: {summary['input_review_notes'] or '[none]'}",
        f"Output DOCX: {summary['output_docx']}",
        f"Output remediation report: {summary['output_report']}",
        f"Remediated DOCX written: {summary['docx_written'] if not DRY_RUN else 'would write'}",
        f"Remediation report written: {summary['report_written'] if not DRY_RUN else 'would write'}",
        f"Slides processed: {summary['slides_processed']}",
        f"Slides with warnings: {summary['slides_with_warnings']}",
        f"Detected visual entries: {summary['detected_visual_entries']}",
        f"Critical warnings carried forward: {summary['critical_warnings']}",
        f"Slides missing images: {summary['slides_missing_images']}",
    ]
    print("\n".join(lines))


def main() -> int:
    setup_logging()

    input_docx = Path(INPUT_DOCX_PATH).expanduser() if INPUT_DOCX_PATH else Path("./input_companion.docx")
    input_manifest = Path(INPUT_MANIFEST_JSON_PATH).expanduser() if INPUT_MANIFEST_JSON_PATH else Path("./input_extract_manifest.json")
    review_notes_path = Path(INPUT_REVIEW_NOTES_PATH).expanduser() if INPUT_REVIEW_NOTES_PATH else None
    output_paths = derive_output_paths(input_docx)

    manifest: dict[str, Any] = {
        "manifest_summary": {
            "total_slides": 0,
            "detected_visual_entries": 0,
        },
        "slide_entries": [],
    }
    image_results: dict[int, ImageResolutionResult] = {}

    try:
        input_docx, input_manifest, review_notes_path, output_paths = validate_config()
        ensure_output_directories(output_paths)

        manifest = load_manifest(input_manifest)
        validate_manifest_against_inputs(manifest, input_docx)
        review_notes_info = read_review_notes(review_notes_path)

        slide_image_dir = resolve_slide_image_directory(manifest)
        for slide in manifest.get("slide_entries", []):
            result = resolve_slide_image(slide, slide_image_dir)
            image_results[slide["slide_number"]] = result

        critical = unresolved_critical_warnings(manifest, image_results)
        if critical and FAIL_ON_UNRESOLVED_CRITICAL_WARNINGS:
            message = "Critical warnings remain unresolved according to the manifest and current configuration blocks output."
            summary = build_run_summary(
                STATUS_FAILURE,
                message,
                input_docx,
                input_manifest,
                review_notes_path,
                output_paths,
                manifest,
                (False, False),
                image_results,
            )
            print_run_summary(summary)
            return fail(message, EXIT_RUNTIME_ERROR)

        document = create_document_from_manifest(manifest, image_results)
        report_text = build_remediation_report(
            manifest,
            review_notes_info,
            input_docx,
            input_manifest,
            output_paths.docx,
            image_results,
        )
        outputs_written = write_outputs(output_paths, document, report_text)

        if DRY_RUN:
            status = STATUS_DRY_RUN_COMPLETE
            message = "Validation, reconstruction, image resolution, and reporting completed. No files were written."
        else:
            status = STATUS_SUCCESS_WITH_WARNINGS if critical or any(
                gather_slide_warnings(slide) or image_results[slide["slide_number"]].warnings
                for slide in manifest.get("slide_entries", [])
            ) else STATUS_SUCCESS
            message = (
                "Remediation completed. Manual review is still required."
                if status == STATUS_SUCCESS_WITH_WARNINGS
                else "Remediation completed without carried-forward warnings."
            )

        summary = build_run_summary(
            status,
            message,
            input_docx,
            input_manifest,
            review_notes_path,
            output_paths,
            manifest,
            outputs_written,
            image_results,
        )
        print_run_summary(summary)

        if status == STATUS_SUCCESS_WITH_WARNINGS:
            LOGGER.warning("Remediation completed with carried-forward warnings.")
        elif status == STATUS_SUCCESS:
            LOGGER.info("Remediation completed cleanly.")
        elif status == STATUS_DRY_RUN_COMPLETE:
            LOGGER.info("Dry run completed.")

        return EXIT_SUCCESS

    except (ConfigError, ManifestError) as exc:
        summary = build_run_summary(
            STATUS_FAILURE,
            str(exc),
            input_docx,
            input_manifest,
            review_notes_path,
            output_paths,
            manifest,
            (False, False),
            image_results,
        )
        print_run_summary(summary)
        return fail(str(exc), EXIT_CONFIG_ERROR)

    except Exception as exc:
        summary = build_run_summary(
            STATUS_FAILURE,
            str(exc),
            input_docx,
            input_manifest,
            review_notes_path,
            output_paths,
            manifest,
            (False, False),
            image_results,
        )
        print_run_summary(summary)
        return fail(str(exc), EXIT_RUNTIME_ERROR)


if __name__ == "__main__":
    sys.exit(main())
